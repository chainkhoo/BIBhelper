# coding=utf-8
import importlib
import logging
import os
import platform
import re
import shutil
import subprocess
import sys
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
from pathlib import Path

SYSTEM = platform.system()
FILE_METADATA = {}

# ==============================================================================
# SECTION 0: 虚拟环境与依赖管理 (保持不变)
# ==============================================================================
VENV_DIR = Path(__file__).resolve().parent / "aiahelper_pro"
VENV_PYTHON = VENV_DIR / ("Scripts/python.exe" if os.name == "nt" else "bin/python")

def in_venv():
    return sys.prefix != sys.base_prefix

if not in_venv():
    if not VENV_DIR.exists():
        print(f"📦 创建虚拟环境: {VENV_DIR}")
        subprocess.check_call([sys.executable, "-m", "venv", str(VENV_DIR)])
    print("📈 升级 pip")
    subprocess.check_call([str(VENV_PYTHON), "-m", "pip", "install", "--upgrade", "pip"])
    print("🚀 重启脚本到虚拟环境执行")
    subprocess.check_call([str(VENV_PYTHON), __file__] + sys.argv[1:])
    sys.exit(0)

def ensure_package(pkg, module_name=None):
    try:
        importlib.import_module(module_name or pkg)
    except ImportError:
        print(f"📦 安装缺失依赖: {pkg}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg])

DEPENDENCIES = [
    ("colorama", "colorama"),
    ("pdfplumber", "pdfplumber"),
    ("requests", "requests"),
    ("python-docx", "docx"),
    ("docx2pdf", "docx2pdf"),
    ("pymupdf", "fitz"),
]
for package_name, module_name in DEPENDENCIES:
    ensure_package(package_name, module_name)

import pdfplumber, requests, fitz
from docx import Document
from docx2pdf import convert
from colorama import Fore, Style, init
init(autoreset=True)

# 抑制 pdfminer 关于缺失 FontBBox 的警告输出
logging.getLogger("pdfminer").setLevel(logging.ERROR)

def print_info(m): print(Fore.BLUE + m + Style.RESET_ALL)
def print_success(m): print(Fore.GREEN + m + Style.RESET_ALL)
def print_warn(m): print(Fore.YELLOW + m + Style.RESET_ALL)
def print_error(m): print(Fore.RED + m + Style.RESET_ALL)

def check_libreoffice():
    """检查LibreOffice是否已安装"""
    # 检查系统PATH中的soffice
    soffice_path = shutil.which("soffice")
    if soffice_path:
        print_success(f"✅ 检测到LibreOffice: {soffice_path}")
        return True
    
    # 在macOS上检查常见安装路径
    if SYSTEM == "Darwin":
        possible_paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/usr/local/bin/soffice",
            "/opt/homebrew/bin/soffice"
        ]
        for path in possible_paths:
            if Path(path).exists():
                print_success(f"✅ 检测到LibreOffice: {path}")
                return True
    
    # 在Windows上检查常见安装路径
    elif SYSTEM == "Windows":
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
        ]
        for path in possible_paths:
            if Path(path).exists():
                print_success(f"✅ 检测到LibreOffice: {path}")
                return True
    
    return False

def install_libreoffice_instructions():
    """提供LibreOffice安装指导"""
    print_warn("⚠️ 未检测到LibreOffice，这是PDF转换的首选工具")
    print_info("📋 LibreOffice安装指导:")
    
    if SYSTEM == "Darwin":
        print_info("  macOS:")
        print_info("    1. 使用Homebrew安装: brew install --cask libreoffice")
        print_info("    2. 或从官网下载: https://www.libreoffice.org/download/download/")
        print_info("    3. 安装后重启终端")
    elif SYSTEM == "Windows":
        print_info("  Windows:")
        print_info("    1. 从官网下载: https://www.libreoffice.org/download/download/")
        print_info("    2. 运行安装程序并添加到系统PATH")
    else:
        print_info("  Linux:")
        print_info("    1. Ubuntu/Debian: sudo apt install libreoffice")
        print_info("    2. CentOS/RHEL: sudo yum install libreoffice")
        print_info("    3. 或从官网下载: https://www.libreoffice.org/download/download/")
    
    print_info("💡 安装完成后重新运行此程序即可使用LibreOffice进行PDF转换")

def check_pdf_conversion_tools():
    """检查PDF转换工具并给出建议"""
    print_info("🔍 检查PDF转换工具...")
    
    tools_available = []
    
    # 检查LibreOffice（首选）
    if check_libreoffice():
        tools_available.append("LibreOffice")
    else:
        install_libreoffice_instructions()
    
    # 检查Pages（仅macOS）
    if SYSTEM == "Darwin":
        try:
            result = subprocess.run(["osascript", "-e", "tell application \"System Events\" to get name of every process"], 
                                  capture_output=True, text=True, timeout=5)
            if "Pages" in result.stdout:
                print_success("✅ 检测到Pages应用")
                tools_available.append("Pages")
            else:
                print_warn("⚠️ 未检测到Pages应用")
        except:
            print_warn("⚠️ 无法检测Pages应用状态")
    
    # 检查Microsoft Word（用于docx2pdf）
    if SYSTEM == "Darwin":
        try:
            result = subprocess.run(["mdfind", "kMDItemCFBundleIdentifier == 'com.microsoft.Word'"], 
                                  capture_output=True, text=True, timeout=5)
            if result.stdout.strip():
                print_success("✅ 检测到Microsoft Word")
                tools_available.append("Microsoft Word")
            else:
                print_warn("⚠️ 未检测到Microsoft Word")
        except:
            print_warn("⚠️ 无法检测Microsoft Word状态")
    
    if tools_available:
        print_success(f"✅ 可用的PDF转换工具: {', '.join(tools_available)}")
        if "LibreOffice" in tools_available:
            print_success("🎯 将优先使用LibreOffice进行PDF转换")
    else:
        print_error("❌ 未检测到任何PDF转换工具")
        print_warn("⚠️ 建议安装LibreOffice以获得最佳PDF转换体验")
    
    return len(tools_available) > 0

# ==============================================================================
# SECTION 1: 全局配置 (新增)
# ==============================================================================
PLAN_CONFIG = {
    'savings': {
        'name': '储蓄险',
        'templates': {
            'single': 'templates/template_savings_standalone.docx',
            'comparison': 'templates/template_savings_comparison.docx'
        }
    },
    'critical_illness': {
        'name': '重疾险',
        'templates': {
            'single': 'templates/template_ci_single.docx'
        }
    }
}

ANNOTATION_OVERLAY_PATH = Path(__file__).resolve().parent / "templates" / "aia_annotation_overlay.png"
DETAIL_SECTION_KEYWORDS = ["详细说明", "詳細說明"]
OVERLAY_SETTINGS = {
    'default': {
        'fit': 'contain',      # contain | width | height | cover
        'anchor': 'center',    # center | top-left | bottom-right ...
        'offset_x': 0,
        'offset_y': 0,
    },
    'savings': {},
}
_OVERLAY_DIMENSIONS = None

# ==============================================================================
# SECTION 2: 文件扫描与任务决策 (核心重构)
# ==============================================================================
def _clean_insured_name(name):
    if not name:
        return None
    name = name.strip()
    for stop in (" 年龄", " 年齡", " 性别", " 性別", "\n年龄", "年龄", "性别", "非吸烟者", "非吸煙者"):
        idx = name.find(stop)
        if idx != -1:
            name = name[:idx]
    name = name.strip()
    for suffix in ("先生", "女士", "小姐", "太太", "宝宝"):
        if name.endswith(suffix):
            name = name[:-len(suffix)]
            break
    return name.strip()


def _parse_to_unicode_map(cmap_stream):
    mapping = {}
    try:
        data = cmap_stream.get_data().decode('latin1')
    except Exception:
        return mapping
    lines = iter(data.splitlines())
    import re
    for line in lines:
        line = line.strip()
        if line.endswith('beginbfchar'):
            count = int(line.split()[0])
            for _ in range(count):
                src, dst = re.findall(r'<([^>]+)>', next(lines))
                try:
                    ch = bytes.fromhex(src).decode('latin1')
                    mapped = bytes.fromhex(dst).decode('utf-16-be')
                    if ch != mapped:
                        mapping[ch] = mapped
                except Exception:
                    continue
        elif line.endswith('beginbfrange'):
            count = int(line.split()[0])
            for _ in range(count):
                parts = re.findall(r'<([^>]+)>', next(lines))
                if len(parts) == 3:
                    start, end, dest = parts
                    start_int = int(start, 16)
                    end_int = int(end, 16)
                    try:
                        if len(dest) == 4:
                            base = int(dest, 16)
                            for offset, code in enumerate(range(start_int, end_int + 1)):
                                ch = bytes([code]).decode('latin1')
                                mapped = chr(base + offset)
                                if ch != mapped:
                                    mapping[ch] = mapped
                        else:
                            dest_bytes = bytes.fromhex(dest)
                            for offset, code in enumerate(range(start_int, end_int + 1)):
                                chunk = dest_bytes[offset*2:(offset+1)*2]
                                mapped = chunk.decode('utf-16-be')
                                ch = bytes([code]).decode('latin1')
                                if ch != mapped:
                                    mapping[ch] = mapped
                    except Exception:
                        continue
                elif len(parts) == 2:
                    start, dest = parts
                    try:
                        ch = bytes.fromhex(start).decode('latin1')
                        mapped = bytes.fromhex(dest).decode('utf-16-be')
                        if ch != mapped:
                            mapping[ch] = mapped
                    except Exception:
                        continue
    return mapping


def _decode_special_sequences(pdf_path, text, pdf_obj=None):
    try:
        import pdfplumber
        from pdfplumber.utils import resolve
    except Exception:
        return text

    pdf = pdf_obj
    close_after = False
    if pdf is None:
        try:
            pdf = pdfplumber.open(pdf_path)
            close_after = True
        except Exception:
            return text

    custom_map = {}
    try:
        for page in pdf.pages[:3]:
            try:
                resources = resolve(page.page_obj.attrs['Resources'])
                fonts = resources.get('Font', {})
                for ref in fonts.values():
                    font = resolve(ref)
                    to_unicode_ref = font.get('ToUnicode')
                    if to_unicode_ref:
                        cmap = resolve(to_unicode_ref)
                        custom_map.update(_parse_to_unicode_map(cmap))
            except Exception:
                continue
    finally:
        if close_after:
            pdf.close()

    custom_map = {k: v for k, v in custom_map.items() if len(k) == 1 and len(v) == 1}
    if not custom_map:
        return text

    import re
    pattern = re.compile('([' + re.escape(''.join(custom_map.keys())) + r'\s]{2,})')

    text = text.replace('’', "'")

    def repl(match):
        raw = match.group(1).replace(' ', '')
        if not raw:
            return match.group(1)
        return ''.join(custom_map.get(ch, ch) for ch in raw)

    return pattern.sub(repl, text)


def extract_payment_term_and_age(pdf_path):
    """从PDF中提取缴费年限、年龄及受保人姓名"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # 扫描前3页
            all_text = ""
            for page in pdf.pages[:3]:
                page_text = page.extract_text() or ""
                all_text += page_text + "\n"

            all_text = _decode_special_sequences(pdf_path, all_text, pdf)

            # 提取缴费年限
            payment_term = None
            payment_patterns = [
                r'(\d+)\s*年[缴繳][费費]',
                r'[缴繳][费費]年期[：:]\s*(\d+)',
                r'供款年期[：:]\s*(\d+)',
                r'[缴繳][费費]期[间間][：:]\s*(\d+)'
            ]
            
            for pattern in payment_patterns:
                matches = re.findall(pattern, all_text)
                if matches:
                    payment_term = int(matches[0])  # 取第一个匹配
                    break
            
            # 提取年龄
            age = None
            age_patterns = [
                r'[年龄齡][：:]\s*(\d+)',
                r'投保[年龄齡][：:]\s*(\d+)',
                r'受保人.*?(\d+)[岁歲]'
            ]
            
            for pattern in age_patterns:
                matches = re.findall(pattern, all_text)
                if matches:
                    age = int(matches[0])  # 取第一个匹配
                    break
            
            # 提取姓名
            name = None
            name_patterns = [
                r'受保人姓名[：:\s]*([\u4e00-\u9fa5A-Za-z]+)',
                r'受保人[：:\s]*([\u4e00-\u9fa5A-Za-z]+)',
                r'受保人姓名[：:\s]*([\u4e00-\u9fa5A-Za-z]+)[\s]*(?:先生|女士|小姐|太太|宝宝|寶寶)',
                r'([\u4e00-\u9fa5]{1,4})(?:先生|女士|小姐|太太|宝宝|寶寶)'
            ]
            for pattern in name_patterns:
                matches = re.findall(pattern, all_text)
                if matches:
                    candidate = _clean_insured_name(matches[0])
                    if candidate:
                        name = candidate
                        break
            
            return payment_term, age, name

    except Exception as e:
        print_warn(f"  - ⚠️ 提取 {pdf_path} 信息时出错: {e}")
        return None, None, None


def _extract_total_premium(text):
    """从文本中提取年缴保费：投保时年缴总保费 - 保险业监管局(IA)保费征费"""
    
    # 提取投保时年缴总保费
    total_premium = None
    total_premium_patterns = [
        r"投保时年缴总保费[：:\s]*([\d,]+(?:\.\d+)?)",
        r"投保時年繳總保費[：:\s]*([\d,]+(?:\.\d+)?)",
        r"投保时年缴总保费（已包括保费征费）[：:\s]*([\d,]+(?:\.\d+)?)",
        r"投保時年繳總保費（已包括保費徵費）[：:\s]*([\d,]+(?:\.\d+)?)"
    ]
    
    for pattern in total_premium_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                total_premium = float(match.group(1).replace(',', ''))
                break
            except ValueError:
                continue
    
    # 提取保险业监管局(IA)保费征费
    ia_fee = None
    ia_fee_patterns = [
        r"保险业监管局\s*\(IA\)\s*保费征费[：:\s]*([\d,]+(?:\.\d+)?)",
        r"保險業監管局\s*\(IA\)\s*保費徵費[：:\s]*([\d,]+(?:\.\d+)?)",
        r"IA\s*保费征费[：:\s]*([\d,]+(?:\.\d+)?)",
        r"IA\s*保費徵費[：:\s]*([\d,]+(?:\.\d+)?)"
    ]
    
    for pattern in ia_fee_patterns:
        match = re.search(pattern, text)
        if match:
            try:
                ia_fee = float(match.group(1).replace(',', ''))
                break
            except ValueError:
                continue
    
    # 计算最终年缴保费：投保时年缴总保费 - IA保费征费
    if total_premium is not None and ia_fee is not None:
        final_premium = total_premium - ia_fee
        print_info(f"💰 保费计算: {total_premium} - {ia_fee} = {final_premium}")
        return int(round(final_premium))  # 返回整数
    elif total_premium is not None:
        # 如果没有找到IA保费征费，直接返回投保时年缴总保费
        print_warn(f"⚠️ 未找到IA保费征费，使用投保时年缴总保费: {total_premium}")
        return int(round(total_premium))
    
    # 如果上述方法都失败，尝试从表格中提取（保持原有逻辑作为备用）
    try:
        # 查找"分红保单销售说明文件"章节
        dividend_section_match = re.search(r'分红保单销售说明文件.*?(?=\n\s*\d+\.|$)', text, re.DOTALL | re.IGNORECASE)
        if dividend_section_match:
            dividend_section = dividend_section_match.group(0)
            
            # 在该章节中查找"2. 保障摘要"
            summary_match = re.search(r'2\.\s*保障摘要.*?(?=\n\s*\d+\.|$)', dividend_section, re.DOTALL | re.IGNORECASE)
            if summary_match:
                summary_section = summary_match.group(0)
                
                # 在保障摘要中查找"(i) 基本保单"
                basic_policy_match = re.search(r'\(i\)\s*基本保单.*?(?=\n\s*\([a-z]+\)|$)', summary_section, re.DOTALL | re.IGNORECASE)
                if basic_policy_match:
                    basic_policy_section = basic_policy_match.group(0)
                    
                    # 将文本按行分割，查找表格
                    lines = basic_policy_section.split('\n')
                    for i, line in enumerate(lines):
                        # 跳过标题行，查找第2行数据
                        if i > 0 and line.strip():  # 跳过第一行（标题）
                            # 提取该行中的所有数字（包括带逗号的）
                            numbers = re.findall(r'\d{1,3}(?:,\d{3})*(?:\.\d{2})?', line)
                            if len(numbers) >= 4:  # 确保有第4列
                                try:
                                    # 取第4列的值，只保留整数部分
                                    value = float(numbers[3].replace(',', ''))
                                    return int(value)  # 只保留整数部分
                                except (ValueError, IndexError):
                                    continue
    except Exception:
        pass
    
    return None

def classify_by_payment_term_and_age(payment_term, age, filename):
    """根据缴费年限和年龄分类保险类型"""

    # 如果无法提取缴费年限，尝试文件名判断
    if payment_term is None:
        # 基于文件名的备用分类
        if '储蓄' in filename or '财富增值' in filename:
            return 'savings'
        elif '重疾' in filename or '疾病保障' in filename:
            return 'critical_illness'
        else:
            return None
    
    if payment_term <= 5:
        # 缴费5年及以下：储蓄险
        return 'savings'
    elif payment_term >= 10:
        # 缴费10年以上：重疾险
        return 'critical_illness'
    else:
        # 其他情况暂时归类为未识别
        return None

def scan_and_classify_pdfs():
    """扫描当前目录下的PDF，并根据缴费年限和年龄进行智能分类。"""
    pdfs = [f for f in os.listdir('.') if f.lower().endswith('.pdf')]
    if not pdfs:
        print_error("❌ 当前目录未找到任何 PDF 文件。")
        sys.exit(1)
        
    print_info("📂 正在扫描和分类目录中的 PDF 文件...")
    classified_pdfs = defaultdict(list)
    file_metadata = {}
    unclassified_pdfs = []
    current_dir = Path.cwd()

    for pdf_path in pdfs:
        abs_path = str((current_dir / pdf_path).resolve())
        try:
            # 提取缴费年限、年龄及姓名
            payment_term, age, name = extract_payment_term_and_age(pdf_path)
            
            # 根据缴费年限、年龄和文件名分类
            plan_type = classify_by_payment_term_and_age(payment_term, age, pdf_path)

            is_segment = False

            if plan_type and plan_type in PLAN_CONFIG:
                classified_pdfs[plan_type].append(pdf_path)
                plan_name = PLAN_CONFIG[plan_type]['name']
                detail = []
                if payment_term:
                    detail.append(f"缴费{payment_term}年")
                if age is not None:
                    detail.append(f"{age}岁")
                detail_text = f" ({', '.join(detail)})" if detail else ""
                print(f"  - {pdf_path} -> {Fore.CYAN}{plan_name}{detail_text}")
            else:
                unclassified_pdfs.append(pdf_path)
                detail = []
                if payment_term:
                    detail.append(f"缴费{payment_term}年")
                if age:
                    detail.append(f"{age}岁")
                detail_text = f" ({', '.join(detail)})" if detail else ""
                print(f"  - {pdf_path} -> {Fore.YELLOW}未识别类型{detail_text}")

            info = {
                'payment_term': payment_term,
                'age': age,
                'name': name,
                'is_segment': is_segment,
                'filename': pdf_path,
                'plan_name': PLAN_CONFIG[plan_type]['name'] if plan_type and plan_type in PLAN_CONFIG else None,
            }
            file_metadata[abs_path] = info
                
        except Exception as e:
            print_warn(f"  - ⚠️ 处理 {pdf_path} 时出错: {e}")
            unclassified_pdfs.append(pdf_path)
            info = {
                'payment_term': None,
                'age': None,
                'name': None,
                'is_segment': False,
                'filename': pdf_path,
                'plan_name': None,
            }
            file_metadata[abs_path] = info
    
    global FILE_METADATA
    FILE_METADATA = file_metadata.copy()

    return classified_pdfs, file_metadata

def _flatten_classified(classified_pdfs):
    flat = []
    for plan_type, files in classified_pdfs.items():
        for f in files:
            flat.append({'type': plan_type, 'file': f})
    return flat

def _print_scan_summary(classified_pdfs):
    total = sum(len(v) for v in classified_pdfs.values())
    print_info(f"🔎 共检测到 {total} 个PDF：")
    for plan_type, files in classified_pdfs.items():
        name = PLAN_CONFIG[plan_type]['name']
        print(f"  - {name}：{len(files)} 个")
    print("")

def _print_execution_suggestion(tasks):
    print_info("🧭 执行建议：")
    for i, t in enumerate(tasks, 1):
        plan_name = PLAN_CONFIG[t['type']]['name']
        template = PLAN_CONFIG[t['type']]['templates'].get(t['mode']) or PLAN_CONFIG[t['type']]['templates'].get('single')
        # 显示文件名而不是完整路径，保持用户友好
        file_list = ", ".join([Path(f).name for f in t['files']])
        mode_name = "对比总结书" if t['mode'] == 'comparison' else "单独总结书"
        print(f"  {i}. {plan_name} - {mode_name} -> 模板: {template} -> 文件: {file_list}")
    print("")


def _deduplicate_tasks(tasks):
    deduped = []
    seen = set()
    for task in tasks:
        key = (task['type'], task['mode'], tuple(sorted(task['files'])))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(task)
    return deduped


def _build_savings_tasks(files, file_metadata):
    grouped = {}
    surname_age_map = {}
    for file_path in files:
        meta = file_metadata.get(file_path, {})
        name = meta.get('name')
        age = meta.get('age')
        if name and len(name) >= 2:
            key = (name[0], age)
            existing = surname_age_map.get(key)
            if not existing or len(name) > len(existing):
                surname_age_map[key] = name

    for file_path in files:
        meta = file_metadata.get(file_path, {})
        raw_name = meta.get('name')
        age = meta.get('age')
        if raw_name and len(raw_name) >= 2:
            key = raw_name
        elif raw_name:
            key = surname_age_map.get((raw_name[0], age)) or f"{raw_name[0]}_{age if age is not None else 'unknown'}"
        elif age is not None:
            key = f"age_{age}_{Path(file_path).stem}"
        else:
            key = Path(file_path).stem

        entry = grouped.setdefault(key, {'files': [], 'name': None})
        entry['files'].append(file_path)
        if raw_name and (entry['name'] is None or len(raw_name) > len(entry['name'])):
            entry['name'] = raw_name

    tasks = []
    for data in grouped.values():
        person_files = sorted(data['files'])
        # 单独总结书
        for path in person_files:
            tasks.append({'type': 'savings', 'mode': 'single', 'files': [path]})
        # 对比总结书（仅使用前两个方案）
        if len(person_files) >= 2:
            tasks.append({'type': 'savings', 'mode': 'comparison', 'files': person_files[:2]})
    return tasks


def _build_critical_tasks(files):
    return [{'type': 'critical_illness', 'mode': 'single', 'files': [path]} for path in files]


def _build_auto_tasks(abs_classified_pdfs, file_metadata):
    tasks = []
    if 'savings' in abs_classified_pdfs:
        tasks.extend(_build_savings_tasks(abs_classified_pdfs['savings'], file_metadata))
    if 'critical_illness' in abs_classified_pdfs:
        tasks.extend(_build_critical_tasks(abs_classified_pdfs['critical_illness']))
    return _deduplicate_tasks(tasks)


def _get_user_confirmation(prompt_text, default_yes=True):
    """统一的用户确认输入处理，默认为是"""
    if default_yes:
        full_prompt = prompt_text + "\n" + Fore.GREEN + "按回车确认执行，输入 n 取消: " + Style.RESET_ALL
        return input(full_prompt).strip().lower() != 'n'
    else:
        full_prompt = prompt_text + "\n" + Fore.YELLOW + "输入 y 确认，按回车跳过: " + Style.RESET_ALL  
        return input(full_prompt).strip().lower() == 'y'

def determine_tasks(classified_pdfs, file_metadata):
    """根据规则生成默认任务建议，并支持手动调整。"""
    tasks = []

    _print_scan_summary(classified_pdfs)

    all_items = _flatten_classified(classified_pdfs)
    current_dir = Path.cwd()

    abs_classified_pdfs = defaultdict(list)
    abs_all_items = []
    for item in all_items:
        abs_path = str(current_dir / item['file'])
        abs_all_items.append({'type': item['type'], 'file': abs_path})
        abs_classified_pdfs[item['type']].append(abs_path)
    all_items = abs_all_items

    auto_tasks = _build_auto_tasks(abs_classified_pdfs, file_metadata)

    if not auto_tasks:
        print_error("❌ 未能根据规则生成任务，请检查输入文件。")
        sys.exit(1)

    print_info("系统建议执行以下任务：")
    _print_execution_suggestion(auto_tasks)
    if _get_user_confirmation("是否按照以上建议执行？"):
        print_success("✅ 已确认执行系统建议的任务")
        return auto_tasks

    # 手动选择模式
    print_warn("进入手动选择模式：")

    # 0) 先选择险种
    available_types = [t for t in PLAN_CONFIG.keys() if abs_classified_pdfs.get(t)]
    print_info("步骤0：先选择要处理的险种")
    for idx, plan_type in enumerate(available_types, 1):
        plan_name = PLAN_CONFIG[plan_type]['name']
        count = len(abs_classified_pdfs.get(plan_type, []))
        print(f"  {idx}. {plan_name}（{count} 个文件）")

    type_input = input(Fore.YELLOW + "请输入险种序号（可多个，空格分隔），直接回车退出: " + Style.RESET_ALL).strip()
    if not type_input:
        print_info("未选择险种。程序退出。")
        sys.exit(0)

    selected_types = []
    for token in type_input.split():
        try:
            idx = int(token)
        except ValueError:
            print_warn(f"险种序号无效：{token}，已跳过。")
            continue
        if idx < 1 or idx > len(available_types):
            print_warn(f"险种序号超出范围：{idx}，已跳过。")
            continue
        selected_types.append(available_types[idx - 1])
    selected_types = list(dict.fromkeys(selected_types))

    if not selected_types:
        print_info("没有有效的险种选择。程序退出。")
        sys.exit(0)

    filtered_items = [item for item in all_items if item['type'] in selected_types]
    if not filtered_items:
        print_error("❌ 所选险种下没有可用文件。")
        sys.exit(1)

    print_info("已筛选文件列表：")
    for idx, item in enumerate(filtered_items, 1):
        name = PLAN_CONFIG[item['type']]['name']
        print(f"  {idx}. [{name}] {item['file']}")

    # 1) 选择要生成"单独总结书"的文件（可多选）
    print_info("步骤1：选择要生成单独总结书的文件")
    choice_single = input(Fore.YELLOW + "请输入文件序号（可多个，空格分隔），直接回车跳过: " + Style.RESET_ALL).strip()
    selected_indexes = set()
    if choice_single:
        for token in choice_single.split():
            try:
                i = int(token)
            except ValueError:
                print_warn(f"单独总结书序号无效：{token}，已跳过。")
                continue
            if i < 1 or i > len(filtered_items):
                print_warn(f"单独总结书序号超出范围：{i}，已跳过。")
                continue
            selected_indexes.add(i)
        for i in sorted(selected_indexes):
            item = filtered_items[i - 1]
            tasks.append({'type': item['type'], 'mode': 'single', 'files': [item['file']]})

    # 2) 选择要生成"对比总结书"的配对（可多对，用逗号分隔形如 1-3, 2-5）
    comparable_types = [t for t in selected_types if PLAN_CONFIG[t]['templates'].get('comparison')]
    if comparable_types:
        print_info("步骤2：选择要生成对比总结书的文件配对（仅同类型允许对比）")
        print("示例：1-3 或 1-3,2-4（多对用逗号分隔）")
        pair_input = input(Fore.YELLOW + "请输入配对，直接回车跳过: " + Style.RESET_ALL).strip()
        if pair_input:
            pairs = re.split(r"[,\uFF0C]", pair_input)
            for pair in pairs:
                pair = pair.strip()
                if not pair:
                    continue
                if '-' not in pair:
                    print_warn(f"配对格式无效：{pair}，已跳过。")
                    continue
                left, right = pair.split('-', 1)
                try:
                    li, ri = int(left), int(right)
                    if li < 1 or li > len(filtered_items) or ri < 1 or ri > len(filtered_items):
                        print_error(f"配对 {pair} 序号超出范围，已跳过。")
                        continue
                    a, b = filtered_items[li - 1], filtered_items[ri - 1]
                    if li == ri:
                        print_warn(f"配对 {pair} 使用了同一个文件，已跳过。")
                        continue
                    if a['type'] != b['type']:
                        print_error(f"配对 {pair} 跨类型（{PLAN_CONFIG[a['type']]['name']} vs {PLAN_CONFIG[b['type']]['name']}），已跳过。")
                        continue
                    if not PLAN_CONFIG[a['type']]['templates'].get('comparison'):
                        print_warn(f"配对 {pair} 的类型 {PLAN_CONFIG[a['type']]['name']} 不支持对比总结书，已跳过。")
                        continue
                    files = [a['file'], b['file']]
                    tasks.append({'type': a['type'], 'mode': 'comparison', 'files': files})
                except ValueError:
                    print_error(f"配对 {pair} 输入无效，已跳过。")
    else:
        print_info("步骤2：所选险种不支持对比总结书，已跳过配对选择。")

    tasks = _deduplicate_tasks(tasks)

    if not tasks:
        print_info("没有选择任何要执行的任务。程序退出。")
        sys.exit(0)

    _print_execution_suggestion(tasks)
    return tasks

# ==============================================================================
# SECTION 3: 数据提取 (模块化重构)
# ==============================================================================
def parse_savings_plan(text, usd_cny, idx, shared_data=None):
    """
    解析储蓄险方案数据
    
    Args:
        text: PDF提取的文本
        usd_cny: 美元兑人民币汇率
        idx: 方案索引 (0=单独模式, 1,2=对比模式)
        shared_data: 共享数据字典
        
    Returns:
        (data_dict, updated_shared_data): 解析的数据和更新的共享数据
    """
    if shared_data is None:
        shared_data = {}
    
    # 提取基本信息（只在第一次或单独模式时提取）
    if idx <= 1:
        # 姓名、年龄 - 允许名字中有空格（如 "VIP 先生"）
        name_age_patterns = [
            r"受保人姓名[:：]\s*(.+?)\s*(?:先生|女士|小姐|太太|宝宝|寶寶)\s*年[龄齡][:：]\s*(\d+)",
            r"受保人[:：]\s*(.+?)\s*(?:先生|女士|小姐|太太|宝宝|寶寶)\s*年[龄齡][:：]\s*(\d+)",
            # 后备模式：无称谓
            r"受保人姓名[:：]\s*([^\s]+)\s+年[龄齡][:：]\s*(\d+)",
            r"受保人[:：]\s*([^\s]+)\s+年[龄齡][:：]\s*(\d+)"
        ]
        for pattern in name_age_patterns:
            m = re.search(pattern, text)
            if m:
                shared_data["name"] = _clean_insured_name(m.group(1))
                shared_data["age"] = int(m.group(2))
                break
            
        # 保险计划 & 供款期限
        m_plan = re.search(r"建[议議][书書]摘要[:：]\s*([^\(（]+)\s*[\(（]([0-9]+)\s*年[缴繳][费費]", text)
        if m_plan:
            shared_data["plan_name"] = m_plan.group(1).strip()
            shared_data["payment_term"] = m_plan.group(2).strip()
            
        # 计算衍生字段
        if "age" in shared_data:
            age = shared_data["age"]
            shared_data["age_plus_6"] = age + 6
            shared_data["age_plus_10"] = age + 10
            shared_data["years_withdraw"] = 95 - age

    name = shared_data.get("name")
    age = shared_data.get("age")
    plan_name = shared_data.get("plan_name")
    payment_term = shared_data.get("payment_term")
    
    # 提取保费信息
    premium_usd = 0
    premium_total = _extract_total_premium(text)
    if premium_total:
        premium_usd = round(premium_total, 2)
    else:
        m_row = None
        # 繁体: 分紅保單銷售說明文件
        m_section = re.search(r"分[红紅]保[单單][销銷]售[说說]明文件(.+?)(?=保障及利益摘要)", text, re.S)
        if m_section:
            part = m_section.group(1)
            # 繁体: 首 12 個月意外身故賠償
            m_row = re.search(r"首\s*12\s*[个個]月意外身故[赔賠][偿償][^\n]*", part)
        if m_row:
            row_text = m_row.group(0)
            nums = re.findall(r"\d{1,3}(?:,\d{3})*", row_text)
            if len(nums) >= 2:
                premium_usd = float(nums[1].replace(',', ''))
    premium_cny = round(premium_usd * usd_cny / 10000, 1)

    # 提取现金提取相关数据
    balance_usd_55 = balance_cny_55 = 0
    balance_usd_65 = balance_cny_65 = 0
    balance_usd_85 = balance_cny_85 = 0
    cashout_usd_age10 = cashout_cny_age10 = 0
    cashout_usd_55 = cashout_cny_55 = 0
    cashout_usd_65 = cashout_cny_65 = 0
    cashout_usd_75 = cashout_cny_75 = 0
    cashout_usd_85 = cashout_cny_85 = 0
    section_b = re.search(r"[详詳][细細][说說]明(.+?)[现現]金提取[举舉]例", text, re.S)
    part_b = section_b.group(1) if section_b else ""
    section_a = re.search(r"[现現]金提取[举舉]例(.+)", text, re.S)
    part_a = section_a.group(1) if section_a else ""
    # 繁体: 現金提取後之退保發還金額
    m_a2 = re.search(r"([现現]金提取[后後]之退保[发發][还還]金[额額].+)", part_a, re.S)
    part_a2 = m_a2.group(1) if m_a2 else ""

    withdraw_usd = 0
    for line in part_a.splitlines():
        line = line.strip()
        if not line or not line[0].isdigit():
            continue
        nums = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", line)
        if len(nums) >= 6:
            try:
                total_val = float(nums[-1].replace(',', ''))
            except ValueError:
                continue
            if total_val > 0:
                withdraw_usd = int(round(total_val))
                break
    withdraw_cny = withdraw_usd * usd_cny

    m55 = re.search(r"^\s*55[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    balance_usd_55 = int(m55.group(1).replace(",", "")) if m55 else 0
    balance_cny_55 = balance_usd_55 * usd_cny
    balance_cny_55_wan = balance_cny_55 / 10000 if balance_cny_55 else 0

    m65 = re.search(r"^\s*65[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    balance_usd_65 = int(m65.group(1).replace(",", "")) if m65 else balance_usd_65
    balance_cny_65 = balance_usd_65 * usd_cny
    balance_cny_65_wan = balance_cny_65 / 10000 if balance_cny_65 else 0

    m75 = re.search(r"^\s*75[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    balance_usd_75 = int(m75.group(1).replace(",", "")) if m75 else 0
    balance_cny_75 = balance_usd_75 * usd_cny
    balance_cny_75_wan = balance_cny_75 / 10000 if balance_cny_75 else 0

    m85 = re.search(r"^\s*85[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    balance_usd_85 = int(m85.group(1).replace(",", "")) if m85 else balance_usd_85
    balance_cny_85 = balance_usd_85 * usd_cny
    balance_cny_85_wan = balance_cny_85 / 10000 if balance_cny_85 else 0

    age_plus_10 = age + 10 if age else 0
    m10 = re.search(fr"^{age_plus_10}\s+\d+\s+(?:[\d,]+\s+){{4}}([\d,]+)", part_b, re.M)
    cashout_usd_age10 = int(m10.group(1).replace(",", "")) if m10 else 0
    cashout_cny_age10 = cashout_usd_age10 * usd_cny

    m55b = re.search(r"^55\s+\d+\s+(?:[\d,]+\s+){4}([\d,]+)", part_b, re.M)
    cashout_usd_55 = int(m55b.group(1).replace(",", "")) if m55b else 0
    cashout_cny_55 = cashout_usd_55 * usd_cny

    m65b = re.search(r"^65\s+\d+\s+(?:[\d,]+\s+){4}([\d,]+)", part_b, re.M)
    cashout_usd_65 = int(m65b.group(1).replace(",", "")) if m65b else balance_usd_65
    cashout_cny_65 = cashout_usd_65 * usd_cny

    m75b = re.search(r"^75\s+\d+\s+(?:[\d,]+\s+){4}([\d,]+)", part_b, re.M)
    cashout_usd_75 = int(m75b.group(1).replace(",", "")) if m75b else 0
    cashout_cny_75 = cashout_usd_75 * usd_cny

    m85b = re.search(r"^85\s+\d+\s+(?:[\d,]+\s+){4}([\d,]+)", part_b, re.M)
    cashout_usd_85 = int(m85b.group(1).replace(",", "")) if m85b else 0
    cashout_cny_85 = cashout_usd_85 * usd_cny

    years_withdraw = 95 - age if age else 0
    age_plus_6 = age + 6 if age else 0

    withdraw_cny_wan = round(withdraw_cny / 10000, 2) if withdraw_cny else 0
    withdraw_cny_month = withdraw_cny / 12 if withdraw_cny else 0
    withdraw_cny_total_wan = withdraw_cny * years_withdraw / 10000 if withdraw_cny else 0

    balance_cny_65_wan = balance_cny_65 / 10000 if balance_cny_65 else 0
    balance_cny_85_wan = balance_cny_85 / 10000 if balance_cny_85 else 0

    cashout_cny_age10_wan = cashout_cny_age10 / 10000 if cashout_cny_age10 else 0
    cashout_cny_55_wan = cashout_cny_55 / 10000 if cashout_cny_55 else 0
    cashout_cny_65_wan = cashout_cny_65 / 10000 if cashout_cny_65 else 0
    cashout_cny_75_wan = cashout_cny_75 / 10000 if cashout_cny_75 else 0
    cashout_cny_85_wan = cashout_cny_85 / 10000 if cashout_cny_85 else 0
    cashout_cny_85_qianwan = round(cashout_cny_85 / 10000000) if cashout_cny_85 else 0

    data = {}
    # 单独模式 (idx=0) 或对比模式的第一个文件 (idx=1) 时添加基本字段
    if idx <= 1:
        data["name"] = name
        data["age"] = age
        data["plan_name"] = plan_name
        data["payment_term"] = payment_term
        data["age_plus_6"] = age_plus_6
        data["age_plus_10"] = age_plus_10
        data["years_withdraw"] = years_withdraw

    data.update({
        f"premium_usd_{idx}": premium_usd, f"premium_cny_{idx}": premium_cny,
        f"withdraw_usd_{idx}": withdraw_usd, f"withdraw_cny_{idx}": withdraw_cny,
        f"withdraw_cny_wan_{idx}": withdraw_cny_wan, f"withdraw_cny_month_{idx}": withdraw_cny_month,
        f"withdraw_cny_total_wan_{idx}": withdraw_cny_total_wan,
        f"balance_usd_55_{idx}": balance_usd_55, f"balance_cny_55_{idx}": balance_cny_55, f"balance_cny_55_wan_{idx}": balance_cny_55_wan,
        f"balance_usd_65_{idx}": balance_usd_65, f"balance_cny_65_{idx}": balance_cny_65, f"balance_cny_65_wan_{idx}": balance_cny_65_wan,
        f"balance_usd_75_{idx}": balance_usd_75, f"balance_cny_75_{idx}": balance_cny_75, f"balance_cny_75_wan_{idx}": balance_cny_75_wan,
        f"balance_usd_85_{idx}": balance_usd_85, f"balance_cny_85_{idx}": balance_cny_85, f"balance_cny_85_wan_{idx}": balance_cny_85_wan,
        f"cashout_usd_age_plus_10_{idx}": cashout_usd_age10, f"cashout_cny_age_plus_10_{idx}": cashout_cny_age10, f"cashout_cny_age_plus_10_wan_{idx}": cashout_cny_age10_wan,
        f"cashout_usd_55_{idx}": cashout_usd_55, f"cashout_cny_55_{idx}": cashout_cny_55, f"cashout_cny_55_wan_{idx}": cashout_cny_55_wan,
        f"cashout_usd_65_{idx}": cashout_usd_65, f"cashout_cny_65_{idx}": cashout_cny_65, f"cashout_cny_65_wan_{idx}": cashout_cny_65_wan,
        f"cashout_usd_75_{idx}": cashout_usd_75, f"cashout_cny_75_{idx}": cashout_cny_75, f"cashout_cny_75_wan_{idx}": cashout_cny_75_wan,
        f"cashout_usd_85_{idx}": cashout_usd_85, f"cashout_cny_85_{idx}": cashout_cny_85, f"cashout_cny_85_wan_{idx}": cashout_cny_85_wan,
        f"cashout_cny_85_qianwan_{idx}": cashout_cny_85_qianwan
    })

    if idx == 0:
        payment_term_val = 0
        if shared_data.get("payment_term"):
            try:
                payment_term_val = int(str(shared_data["payment_term"]))
            except ValueError:
                payment_term_val = 0
        premium_usd_all = round(premium_usd * payment_term_val, 2) if premium_usd and payment_term_val else 0
        data["premium_usd_all"] = premium_usd_all
        premium_cny_all_wan = (premium_usd_all * usd_cny) / 10000 if premium_usd_all > 0 and usd_cny > 0 else 0
        if premium_cny_all_wan > 0 and withdraw_cny_total_wan > 0:
            data["withdraw_multiple"] = round(withdraw_cny_total_wan / premium_cny_all_wan, 1)
        else:
            data["withdraw_multiple"] = 0

    keep_decimal_fields = ["withdraw_cny_wan", "premium_cny", "premium_usd", "premium_usd_all", "withdraw_multiple"]
    for k, v in data.items():
        if isinstance(v, float) and not any(k.startswith(f) for f in keep_decimal_fields):
            try: data[k] = int(v)
            except: pass

    if name:
        shared_data["name"] = name
    if age is not None:
        shared_data["age"] = age
    if plan_name:
        shared_data["plan_name"] = plan_name
    if payment_term:
        shared_data["payment_term"] = payment_term

    return data, shared_data

def parse_critical_illness_plan(text, usd_cny, idx, shared_data=None):
    """
    解析重疾险方案数据
    
    Args:
        text: PDF提取的文本
        usd_cny: 美元兑人民币汇率
        idx: 方案索引 (0=单独模式)
        shared_data: 共享数据字典
        
    Returns:
        (data_dict, updated_shared_data): 解析的数据和更新的共享数据
    """
    if shared_data is None:
        shared_data = {}
    
    # 提取基本信息（参照储蓄险逻辑）- 允许名字中有空格
    name_age_patterns = [
        r"受保人姓名[:：]\s*(.+?)\s*(?:先生|女士|小姐|太太|宝宝|寶寶)\s*年[龄齡][:：]\s*(\d+)",
        r"受保人[:：]\s*(.+?)\s*(?:先生|女士|小姐|太太|宝宝|寶寶)\s*年[龄齡][:：]\s*(\d+)",
        # 后备模式：无称谓
        r"受保人姓名[:：]\s*([^\s]+)\s+年[龄齡][:：]\s*(\d+)",
        r"受保人[:：]\s*([^\s]+)\s+年[龄齡][:：]\s*(\d+)"
    ]
    for pattern in name_age_patterns:
        m = re.search(pattern, text)
        if m:
            shared_data["name"] = _clean_insured_name(m.group(1))
            shared_data["age"] = int(m.group(2))
            break

    age = shared_data.get("age")
    if age is not None:
        shared_data["age_plus_10"] = age + 10

    # 判断性别（先生/女士）
    if "先生" in text:
        shared_data["gender"] = "男"
    elif "女士" in text or "小姐" in text or "太太" in text:
        shared_data["gender"] = "女"

    # 提取吸烟信息
    # 优先匹配"非吸烟者"（避免误判）
    if re.search(r"非吸烟者|非吸煙者", text):
        shared_data["smoke"] = "非吸烟者"
    # 匹配"吸烟者"（直接出现，不包含"非"）
    elif re.search(r"吸烟者|吸煙者", text):
        shared_data["smoke"] = "吸烟者"
    # 兼容"吸烟/吸煙："后面跟着"是"或"Yes"的格式
    else:
        smoke_match = re.search(r"[吸烟吸煙][:：]\s*([^\n]+)", text)
        if smoke_match:
            smoke_text = smoke_match.group(1).strip()
            if "是" in smoke_text or "Yes" in smoke_text:
                shared_data["smoke"] = "吸烟者"
            else:
                shared_data["smoke"] = "非吸烟者"
        else:
            shared_data["smoke"] = "非吸烟者"
    
    # 提取保险计划信息
    m_plan = re.search(r"建[议議][书書]摘要[:：]\s*([^\(（]+)\s*[\(（]([0-9]+)\s*年[缴繳][费費]", text)
    if m_plan:
        shared_data["plan_name"] = m_plan.group(1).strip()
        shared_data["payment_term"] = m_plan.group(2).strip()
    
    # 提取保费信息（参照储蓄险逻辑）
    premium_usd = 0
    # 繁体: 分紅保單銷售說明文件
    m_section = re.search(r"分[红紅]保[单單][销銷]售[说說]明文件(.+?)(?=保障及利益摘要)", text, re.S)
    premium_total = _extract_total_premium(text)
    if premium_total:
        premium_usd = round(premium_total, 2)
    elif m_section:
        part = m_section.group(1)
        
        # 查找包含"爱伴航"/"愛伴航"的行，提取年缴保费（第三列数字）
        aibanghang_lines = re.findall(r"[^\n]*[爱愛]伴航[^\n]*", part)
        for line in aibanghang_lines:
            # 查找数字，第三列是年缴保费
            nums = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d{2})?", line)
            if len(nums) >= 3:
                try:
                    premium_usd = float(nums[2].replace(',', ''))
                    break
                except:
                    continue
    
    premium_cny = round(premium_usd * usd_cny, 1)  # 元
    
    # 提取保额信息
    coverage_usd = 0  # 基础保额
    coverage_plus_usd = 0  # 赠送保额
    
    if m_section:
        part = m_section.group(1)
        
        # 查找包含产品名"爱伴航/愛伴航"的行，提取投保时保额
        aibanghang_lines = re.findall(r"[^\n]*[爱愛]伴航[^\n]*", part)
        if aibanghang_lines:
            for line in aibanghang_lines:
                numbers = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", line)
                if not numbers:
                    continue
                try:
                    values = [float(n.replace(',', '')) for n in numbers]
                except:
                    continue
                candidate = max(values)
                if candidate >= 1000:
                    coverage_usd = int(candidate)
                    break
        
        # 如果没找到爱伴航，查找基本保单的投保时保额 (繁体: 基本保單)
        if coverage_usd == 0:
            basic_lines = re.findall(r"[^\n]*基本保[单單][^\n]*", part)
            if basic_lines:
                for line in basic_lines:
                    numbers = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", line)
                    if not numbers:
                        continue
                    try:
                        values = [float(n.replace(',', '')) for n in numbers]
                    except:
                        continue
                    candidate = max(values)
                    if candidate >= 1000:
                        coverage_usd = int(candidate)
                        break

        # 查找升级保障的投保时保额 (繁体: 升級保障)
        upgrade_lines = re.findall(r"[^\n]*升[级級]保障[^\n]*", part)
        if upgrade_lines:
            for line in upgrade_lines:
                numbers = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", line)
                if not numbers:
                    continue
                try:
                    values = [float(n.replace(',', '')) for n in numbers]
                except:
                    continue
                candidate = max(values)
                if candidate >= 1000:
                    coverage_plus_usd = int(candidate)
                    break
    
    # 计算总保额（CNY字段都转换为万元）
    coverage_total_usd = coverage_usd + coverage_plus_usd
    coverage_cny = round(coverage_usd * usd_cny / 10000, 1)  # 万元
    coverage_plus_cny = round(coverage_plus_usd * usd_cny / 10000, 1)  # 万元
    coverage_total_cny = round(coverage_total_usd * usd_cny / 10000, 1)  # 万元
    
    # 提取各年龄段保额和退保金额
    age_plus_10 = shared_data.get("age_plus_10", 0)
    
    # 从"详细说明"章节的表格中提取数据
    coverage_usd_p10 = 0
    coverage_usd_65 = 0
    coverage_usd_85 = 0
    cashout_usd_65 = 0
    cashout_usd_75 = 0
    cashout_usd_85 = 0
    
    # 查找"详细说明"章节 (简体+繁体)
    detail_section = re.search(r"[详詳][细細][说說]明(.+?)(?=保障及利益摘要|$)", text, re.S)
    if detail_section:
        detail_part = detail_section.group(1)
        
        def _find_detail_row(label):
            patterns = [
                rf"^\s*{label}\s+[^\n]*",
                rf"^\s*{label}岁\s+[^\n]*"
            ]
            for pattern in patterns:
                match = re.search(pattern, detail_part, re.M)
                if match:
                    return match.group(0)
            return None

        def _parse_row_values(label):
            line = _find_detail_row(label)
            if not line:
                return None
            nums = re.findall(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?", line)
            if len(nums) < 9:
                return None
            values = [float(n.replace(',', '')) for n in nums]
            age_val = int(round(values[0]))
            policy_year = int(round(values[1]))
            data_vals = [int(round(v)) for v in values[2:]]
            if len(data_vals) < 7:
                return None
            return {
                'age': age_val,
                'policy_year': policy_year,
                'premium_total': data_vals[0],
                'surrender_guaranteed': data_vals[1],
                'surrender_bonus': data_vals[2],
                'surrender_total': data_vals[3],
                'coverage_guaranteed': data_vals[4],
                'coverage_bonus': data_vals[5],
                'coverage_total': data_vals[6]
            }

        # 查找各年龄段数据
        row_p10 = _parse_row_values(str(age_plus_10)) if age_plus_10 else None
        row_55 = _parse_row_values('55')
        row_65 = _parse_row_values('65')
        row_75 = _parse_row_values('75')
        row_85 = _parse_row_values('85')

        if row_p10:
            coverage_usd_p10 = row_p10['coverage_total']
            cashout_usd_age10 = row_p10['surrender_total']

        if row_55:
            cashout_usd_55 = row_55['surrender_total']

        if row_65:
            coverage_usd_65 = row_65['coverage_total']
            cashout_usd_65 = row_65['surrender_total']

        if row_75:
            cashout_usd_75 = row_75['surrender_total']

        if row_85:
            coverage_usd_85 = row_85['coverage_total']
            cashout_usd_85 = row_85['surrender_total']

    # 计算人民币金额（都转换为万元）
    coverage_cny_p10 = round(coverage_usd_p10 * usd_cny / 10000, 1) if coverage_usd_p10 else 0
    coverage_cny_65 = round(coverage_usd_65 * usd_cny / 10000, 1) if coverage_usd_65 else 0
    coverage_cny_85 = round(coverage_usd_85 * usd_cny / 10000, 1) if coverage_usd_85 else 0
    cashout_cny_65 = round(cashout_usd_65 * usd_cny / 10000, 1) if cashout_usd_65 else 0
    cashout_cny_75 = round(cashout_usd_75 * usd_cny / 10000, 1) if cashout_usd_75 else 0
    cashout_cny_85 = round(cashout_usd_85 * usd_cny / 10000, 1) if cashout_usd_85 else 0
    
    # 计算总保费
    payment_term = int(shared_data.get("payment_term", 0)) if shared_data.get("payment_term") else 0
    premium_usd_all = premium_usd * payment_term if premium_usd and payment_term else 0
    if isinstance(premium_usd_all, float):
        premium_usd_all = round(premium_usd_all, 2)
    premium_cny_all_wan = round((premium_usd_all * usd_cny) / 10000, 1) if premium_usd_all > 0 and usd_cny > 0 else 0  # 万元
    
    # 生成数据字典
    data = {
        # 基本信息
        "premium_usd_0": premium_usd,
        "premium_cny_0": premium_cny,
        "premium_usd_all": premium_usd_all,
        "premium_cny_all_wan": premium_cny_all_wan,
        
        # 保额信息
        "coverage_usd": coverage_usd,
        "coverage_cny": coverage_cny,
        "coverage_plus_usd": coverage_plus_usd,
        "coverage_plus_cny": coverage_plus_cny,
        "coverage_total_usd": coverage_total_usd,
        "coverage_total_cny": coverage_total_cny,
        
        # 各年龄段保额
        "coverage_usd_p10": coverage_usd_p10,
        "coverage_cny_p10": coverage_cny_p10,
        "coverage_usd_65": coverage_usd_65,
        "coverage_cny_65": coverage_cny_65,
        "coverage_usd_85": coverage_usd_85,
        "coverage_cny_85": coverage_cny_85,
        
        # 退保金额
        "cashout_usd_65": cashout_usd_65,
        "cashout_cny_65": cashout_cny_65,
        "cashout_usd_75": cashout_usd_75,
        "cashout_cny_75": cashout_cny_75,
        "cashout_usd_85": cashout_usd_85,
        "cashout_cny_85": cashout_cny_85,
    }
    
    # 格式化数值
    keep_decimal_fields = ["premium_usd", "premium_cny"]
    for k, v in data.items():
        if isinstance(v, float) and not any(k.startswith(f) for f in keep_decimal_fields):
            try: 
                data[k] = int(v)
            except: 
                pass
    
    return data, shared_data

def parse_education_plan_dual(segment_text, complete_text, usd_cny, shared_data=None):
    """
    解析教育金双文件方案数据
    
    Args:
        segment_text: 分段提取PDF的文本 (Plan A数据源)
        complete_text: 完整方案PDF的文本 (Plan B/C数据源)
        usd_cny: 美元兑人民币汇率
        shared_data: 共享数据字典
        
    Returns:
        (data_dict, updated_shared_data): 解析的数据和更新的共享数据
    """
    if shared_data is None:
        shared_data = {}
    
    # 从完整方案中提取基本信息
    text = complete_text
    
    # 姓名、年龄（支持繁体中文）
    patterns = [
        r"受保人姓名[:：]\s*([^\s]+)\s*(?:先生|女士|寶寶|宝宝)",
        r"[年龄齡][:：]\s*(\d+)",
        r"年齡[:：]\s*(\d+)"  # 繁体
    ]
    
    name_pattern = patterns[0] + ".*?" + patterns[2]
    m = re.search(name_pattern, text, re.S)
    if not m:
        # 尝试繁体
        name_pattern = patterns[1] + ".*?" + patterns[3]
        m = re.search(name_pattern, text, re.S)
    
    if m:
        shared_data["name"] = m.group(1).strip()
        shared_data["age"] = int(m.group(2))
        
    # 保险计划 & 供款期限（支持繁体）
    plan_patterns = [
        r"建[议議][书書]摘要[:：]\s*([^\(（]+)\s*[\(（]([0-9]+)\s*年[缴繳][费費]",
        r"建議書摘要[:：]\s*([^\(（]+)\s*[\(（]([0-9]+)\s*年繳費"  # 繁体
    ]
    
    for pattern in plan_patterns:
        m_plan = re.search(pattern, text)
        if m_plan:
            shared_data["plan_name"] = m_plan.group(1).strip()
            shared_data["payment_term"] = m_plan.group(2).strip()
            break
            
    # 计算衍生字段
    if "age" in shared_data:
        age = shared_data["age"]
        shared_data["age_plus_5"] = age + 5
        shared_data["age_plus_10"] = age + 10
        # 修正的年限计算
        shared_data["years_withdraw_a"] = 19 - (age + 5)  # 19 - age_plus_5
        shared_data["years_withdraw_b"] = 101 - (age + 5)  # 101 - age_plus_5
    
    # 提取保费信息（从完整方案）
    premium_usd = 0
    premium_total = _extract_total_premium(text)
    if premium_total:
        premium_usd = int(round(premium_total))
    else:
        m_row = None
        premium_patterns = [
            # 繁体: 首 12 個月意外身故賠償
            r"首\s*12\s*[个個]月意外身故[赔賠][偿償][^\n]*",
            r"首\s*12\s*個月意外身故賠償[^\n]*"  # 繁体
        ]
        
        for pattern in premium_patterns:
            # 繁体: 分紅保單銷售說明文件
            m_section = re.search(r"分[红紅]保[单單][销銷]售[说說]明文件(.+?)(?=保障及利益摘要)", text, re.S)
            
            if m_section:
                part = m_section.group(1)
                m_row = re.search(pattern, part)
                if m_row:
                    break
        
        if m_row:
            row_text = m_row.group(0)
            nums = re.findall(r"\d{1,3}(?:,\d{3})*", row_text)
            if len(nums) >= 2:
                premium_usd = int(nums[1].replace(',', ''))
    
    premium_cny = round(premium_usd * usd_cny / 10000, 1)
    
    return _extract_education_data(segment_text, complete_text, premium_usd, premium_cny, usd_cny, shared_data)

def _extract_education_data(segment_text, complete_text, premium_usd, premium_cny, usd_cny, shared_data):
    """提取教育金的详细数据"""
    
    # Plan A: 从分段文件提取
    plan_a_data = _extract_plan_a_from_segment(segment_text, usd_cny, shared_data)
    
    # Plan B/C: 从完整文件提取
    plan_bc_data = _extract_plan_bc_from_complete(complete_text, usd_cny, shared_data, premium_usd)
    
    # 合并所有数据
    data = {
        "premium_usd": premium_usd,
        "premium_cny": premium_cny,
    }
    data.update(plan_a_data)
    data.update(plan_bc_data)
    
    # 添加模板兼容的占位符
    data["age_plus_6"] = shared_data.get("age_plus_5", 6) + 1  # 模板中的age_plus_6 = age_plus_5 + 1
    
    # 重新计算withdraw_multiple（使用正确的公式）
    if "_premium_cny_for_multiple" in data and "_payment_term_for_multiple" in data:
        premium_cny = data["_premium_cny_for_multiple"]
        payment_term = data["_payment_term_for_multiple"]
        withdraw_total_a = data.get("withdraw_cny_total_wan_a", 0)
        
        if premium_cny > 0 and payment_term > 0 and withdraw_total_a > 0:
            data["withdraw_multiple"] = round(withdraw_total_a / (premium_cny * payment_term), 1)
        
        # 清理临时变量
        del data["_premium_cny_for_multiple"]
        del data["_payment_term_for_multiple"]
    
    # 格式化数值
    keep_decimal_fields = ["withdraw_cny_wan", "premium_cny", "cashout_cny_75_yi", "withdraw_multiple", "premium_usd"]
    for k, v in data.items():
        if isinstance(v, float) and not any(k.startswith(f) or f in k for f in keep_decimal_fields):
            try: 
                data[k] = int(v)
            except: 
                pass
    
    return data, shared_data

def _extract_plan_a_from_segment(text, usd_cny, shared_data):
    """Plan A: 严格从分段文件的现金提取举例章节提取数据"""
    data = {}
    
    age_plus_5 = shared_data.get("age_plus_5", 6)
    years_withdraw_a = shared_data.get("years_withdraw_a", 0)
    
    # 严格从现金提取举例章节提取（支持繁体字）
    section_a = re.search(r"[现現]金提取[举舉]例(.+)", text, re.S)
    if not section_a:
        section_a = re.search(r"現金提取舉例(.+)", text, re.S)
    
    if not section_a:
        print_warn("分段文件中未找到现金提取举例章节")
        return {}
    
    part_a = section_a.group(1)
    
    # 提取不同年龄段的提取金额
    # a_1: 从age_plus_5年（6岁）的第5年提取（因为6岁没有第6年数据）
    pattern_a1 = fr"^\s*{age_plus_5}\s+5\s+(?:[\d,]+\s+){{3}}([\d,]+)"
    m_a1 = re.search(pattern_a1, part_a, re.M)
    withdraw_usd_a1 = int(m_a1.group(1).replace(",", "")) if m_a1 else 0
    
    # 如果6岁5年是0，尝试查找7岁6年的数据（原来的逻辑）
    if withdraw_usd_a1 == 0:
        pattern_backup = fr"^\s*7\s+6\s+(?:[\d,]+\s+){{3}}([\d,]+)"
        m_backup = re.search(pattern_backup, part_a, re.M)
        withdraw_usd_a1 = int(m_backup.group(1).replace(",", "")) if m_backup else 0
    
    # a_2: 从25岁提取
    pattern_a2 = fr"^\s*25\s+24\s+(?:[\d,]+\s+){{3}}([\d,]+)"
    m_a2_extract = re.search(pattern_a2, part_a, re.M)
    withdraw_usd_a2 = int(m_a2_extract.group(1).replace(",", "")) if m_a2_extract else 0
    
    # a_3: 从55岁提取
    pattern_a3 = fr"^\s*55\s+54\s+(?:[\d,]+\s+){{3}}([\d,]+)"
    m_a3 = re.search(pattern_a3, part_a, re.M)
    withdraw_usd_a3 = int(m_a3.group(1).replace(",", "")) if m_a3 else 0
    
    # 计算对应的人民币金额
    withdraw_cny_a1 = withdraw_usd_a1 * usd_cny
    withdraw_cny_a2 = withdraw_usd_a2 * usd_cny
    withdraw_cny_a3 = withdraw_usd_a3 * usd_cny
    
    withdraw_cny_wan_a1 = round(withdraw_cny_a1 / 10000, 2)
    withdraw_cny_wan_a2 = round(withdraw_cny_a2 / 10000, 2)
    withdraw_cny_wan_a3 = round(withdraw_cny_a3 / 10000, 2)
    
    withdraw_cny_month_a1 = withdraw_cny_a1 / 12
    withdraw_cny_month_a2 = withdraw_cny_a2 / 12
    withdraw_cny_month_a3 = withdraw_cny_a3 / 12
    
    withdraw_cny_total_wan_a1 = withdraw_cny_a1 * years_withdraw_a / 10000 if years_withdraw_a > 0 else 0
    withdraw_cny_total_wan_a2 = withdraw_cny_a2 * 21 / 10000 if years_withdraw_a > 0 else 0
    withdraw_cny_total_wan_a3 = withdraw_cny_a3 * 21 / 10000 if years_withdraw_a > 0 else 0
    
    # 计算总的withdraw_cny_total_wan_a（可能是三个的总和）
    withdraw_cny_total_wan_a = (withdraw_cny_total_wan_a1 + withdraw_cny_total_wan_a2 + withdraw_cny_total_wan_a3)
    
    # 提取余额数据 - 从现金提取后之退保发还金额部分
    m_a2 = re.search(r"([现現]金提取[后後]之退保[发發][还還]金[额額].+)", part_a, re.S)
    if not m_a2:
        m_a2 = re.search(r"(現金提取後之退保發還金額.+)", part_a, re.S)
    
    part_a2 = m_a2.group(1) if m_a2 else ""
    
    # 提取35岁和55岁余额（参考储蓄险的65岁逻辑）
    # 35岁/歲余额
    m35 = re.search(r"^\s*35[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    if not m35:
        m35 = re.search(r"^\s*35歲?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    
    balance_usd_35 = int(m35.group(1).replace(",", "")) if m35 else 0
    balance_cny_35_wan = balance_usd_35 * usd_cny / 10000 if balance_usd_35 else 0
    
    # 55岁/歲余额（Plan A需要的balance_usd_55_1）
    m55 = re.search(r"^\s*55[岁歲]?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    if not m55:
        m55 = re.search(r"^\s*55歲?\s+[^\n]*?([\d,]+)\s*$", part_a2, re.M)
    
    balance_usd_55 = int(m55.group(1).replace(",", "")) if m55 else 0
    balance_cny_55_wan = balance_usd_55 * usd_cny / 10000 if balance_usd_55 else 0
    
    # Plan A完整数据（不同年龄段的不同数据）
    data.update({
        # a_1: 从age_plus_5（6岁）提取
        "withdraw_usd_a_1": withdraw_usd_a1,
        "withdraw_cny_wan_a_1": withdraw_cny_wan_a1,
        "withdraw_cny_month_a_1": withdraw_cny_month_a1,
        "withdraw_cny_total_wan_a_1": withdraw_cny_total_wan_a1,
        # a_2: 从25岁提取
        "withdraw_usd_a_2": withdraw_usd_a2,
        "withdraw_cny_wan_a_2": withdraw_cny_wan_a2,
        "withdraw_cny_month_a_2": withdraw_cny_month_a2,
        "withdraw_cny_total_wan_a_2": withdraw_cny_total_wan_a2,
        # a_3: 从55岁提取
        "withdraw_usd_a_3": withdraw_usd_a3,
        "withdraw_cny_wan_a_3": withdraw_cny_wan_a3,
        "withdraw_cny_month_a_3": withdraw_cny_month_a3,
        "withdraw_cny_total_wan_a_3": withdraw_cny_total_wan_a3,
        # 总和数据
        "withdraw_cny_total_wan_a": withdraw_cny_total_wan_a,
        "years_withdraw_a": years_withdraw_a,
        # 倍数计算（暂时设为0，将在Plan C中计算）
        "withdraw_multiple": 0,
        # 余额数据（从分段方案现金提取举例章节）
        "balance_usd_35_1": balance_usd_35,
        "balance_cny_35_wan_1": balance_cny_35_wan,
        "balance_usd_55_1": balance_usd_55,
        "balance_cny_55_wan_1": balance_cny_55_wan,
    })
    
    return data

def _extract_plan_bc_from_complete(text, usd_cny, shared_data, premium_usd):
    """Plan B: 从完整文件现金提取举例章节提取，Plan C: 从详细说明章节提取"""
    data = {}
    
    years_withdraw_b = shared_data.get("years_withdraw_b", 0)
    age_plus_6 = shared_data.get("age_plus_5", 6) + 1  # Plan B使用age_plus_6 = age_plus_5 + 1 = 7
    
    # ===== Plan B: 严格从现金提取举例章节提取 (简体+繁体) =====
    section_a = re.search(r"[现現]金提取[举舉]例(.+)", text, re.S)
    if not section_a:
        print_warn("完整文件中未找到现金提取举例章节")
        return {}
    
    part_a = section_a.group(1)
    
    # 提取age_plus_6年的提取金额（Plan B从age_plus_6=7岁第6年提取）
    pattern = fr"^\s*{age_plus_6}\s+6\s+(?:[\d,]+\s+){{3}}([\d,]+)"
    m_wd = re.search(pattern, part_a, re.M)
    withdraw_usd_b = int(m_wd.group(1).replace(",", "")) if m_wd else 0
    
    withdraw_cny_b = withdraw_usd_b * usd_cny
    withdraw_cny_wan_b = round(withdraw_cny_b / 10000, 2)
    withdraw_cny_month_b_1 = withdraw_cny_b / 12  # 月度金额
    withdraw_cny_total_wan_b = withdraw_cny_b * years_withdraw_b / 10000 if years_withdraw_b > 0 else 0
    
    # 从现金提取后之退保发还金额部分提取余额数据（参考储蓄险逻辑）
    m_a2 = re.search(r"([现現]金提取[后後]之退保[发發][还還]金[额額].+)", part_a, re.S)
    part_a2 = m_a2.group(1) if m_a2 else ""
    
    # 提取各年龄余额数据（参考储蓄险逻辑）
    balance_ages = [18, 35, 45, 55, 65]  # 包含55岁，因为55_1在Plan B中需要
    balance_data = {}
    
    for balance_age in balance_ages:
        pattern = fr"^\s*{balance_age}岁?\s+[^\n]*?([\d,]+)\s*$"
        m = re.search(pattern, part_a2, re.M)
        balance_usd = int(m.group(1).replace(",", "")) if m else 0
        balance_data[balance_age] = {
            'usd': balance_usd,
            'cny': balance_usd * usd_cny,
            'wan': balance_usd * usd_cny / 10000
        }
    
    # ===== Plan C: 严格从详细说明章节提取（参考储蓄险逻辑）(简体+繁体) =====
    section_b = re.search(r"[详詳][细細][说說]明(.+?)[现現]金提取[举舉]例", text, re.S)
    part_b = section_b.group(1) if section_b else ""
    
    # 提取各年龄段退保金额（参考储蓄险逻辑）
    cashout_ages = [18, 25, 45, 65, 75]
    cashout_data = {}
    
    for cashout_age in cashout_ages:
        if cashout_age == 75:
            # 75岁特殊处理
            pattern = fr"^75\s+\d+\s+(?:[\d,]+\s+){{4}}([\d,]+)"
            m = re.search(pattern, part_b, re.M)
            cashout_usd = int(m.group(1).replace(",", "")) if m else 0
        else:
            # 其他年龄（18, 25, 45, 65）
            pattern = fr"^{cashout_age}\s+\d+\s+(?:[\d,]+\s+){{4}}([\d,]+)"
            m = re.search(pattern, part_b, re.M)
            cashout_usd = int(m.group(1).replace(",", "")) if m else 0
        
        cashout_data[cashout_age] = {
            'usd': cashout_usd,
            'cny': cashout_usd * usd_cny,
            'wan': cashout_usd * usd_cny / 10000
        }
    
    # Plan B完整数据（从完整文件现金提取举例章节）
    data.update({
        # 基础提取数据
        "withdraw_usd_b": withdraw_usd_b,
        "withdraw_cny_wan_b": withdraw_cny_wan_b,
        "withdraw_cny_month_b_1": withdraw_cny_month_b_1,
        "withdraw_cny_total_wan_b": withdraw_cny_total_wan_b,
        "years_withdraw_b": years_withdraw_b,
        # 余额数据
        "balance_usd_18_2": balance_data[18]['usd'],
        "balance_cny_18_wan_2": balance_data[18]['wan'],
        "balance_usd_35_2": balance_data[35]['usd'],
        "balance_cny_35_wan_2": balance_data[35]['wan'],
        "balance_usd_45_2": balance_data[45]['usd'],
        "balance_cny_45_wan_2": balance_data[45]['wan'],
        "balance_usd_65_2": balance_data[65]['usd'],
        "balance_cny_65_wan_2": balance_data[65]['wan'],
        # 注意：balance_usd_55_1和balance_cny_55_wan_1在Plan A中生成
    })
    
    # Plan C完整数据（从详细说明章节）
    data.update({
        # 退保金额数据
        "cashout_usd_18": cashout_data[18]['usd'],
        "cashout_cny_18_wan": cashout_data[18]['wan'],
        "cashout_usd_25": cashout_data[25]['usd'],
        "cashout_cny_25_wan": cashout_data[25]['wan'],
        "cashout_usd_45": cashout_data[45]['usd'],
        "cashout_cny_45_wan": cashout_data[45]['wan'],
        "cashout_usd_65": cashout_data[65]['usd'],
        "cashout_cny_65_wan": cashout_data[65]['wan'],
        # 75岁特殊处理（亿元计算）
        "cashout_cny_75_yi": round(cashout_data[75]['cny'] / 100000000, 2) if cashout_data[75]['cny'] else 0,
    })
    
    # 倍数计算（Plan A中的withdraw_multiple）
    # 修正公式：withdraw_multiple = withdraw_cny_total_wan_a / (premium_cny * payment_term)
    premium_cny = premium_usd * usd_cny / 10000 if premium_usd > 0 else 0
    payment_term = int(shared_data.get("payment_term", 0)) if shared_data.get("payment_term") else 0
    
    if premium_cny > 0 and payment_term > 0:
        # 需要从Plan A数据中获取withdraw_cny_total_wan_a
        # 这个值将在数据合并后计算
        data["_premium_cny_for_multiple"] = premium_cny
        data["_payment_term_for_multiple"] = payment_term
        data["withdraw_multiple"] = 0  # 临时设为0，将在合并后重新计算
    else:
        data["withdraw_multiple"] = 0
    
    return data

def execute_education_task(files, template_path, plan_name, usd_cny, enable_pdf):
    """特殊处理教育金任务：需要两个PDF文件"""
    
    # 验证文件数量
    if len(files) < 2:
        print_error(f"❌ 教育金需要两个PDF文件（分段+完整），当前只有{len(files)}个")
        return False
    
    # 识别哪个是分段文件，哪个是完整文件
    segment_file = None
    complete_file = None
    
    for file_path in files:
        if not Path(file_path).exists():
            print_error(f"❌ 文件不存在: {file_path}")
            return False
            
        filename = Path(file_path).name
        if '分段' in filename:
            segment_file = file_path
        else:
            complete_file = file_path
    
    # 如果没有明确的分段文件，使用第一个作为分段，第二个作为完整
    if not segment_file:
        segment_file = files[0]
        complete_file = files[1]
        print_warn("⚠️ 未检测到明确的分段文件，使用第一个文件作为分段文件")
    elif not complete_file:
        complete_file = files[1] if files[1] != segment_file else files[0]
    
    print_info(f"  - 分段方案文件: {Path(segment_file).name}")
    print_info(f"  - 完整方案文件: {Path(complete_file).name}")
    
    try:
        # 提取文本
        segment_text = extract_text(segment_file)
        complete_text = extract_text(complete_file)
        
        # 使用双文件解析函数
        all_data, shared_data = parse_education_plan_dual(segment_text, complete_text, usd_cny, {})
        
        # 合并共享数据
        all_data.update(shared_data)
        
        # 使用新的文件保存方式（创建投保人姓名文件夹）
        save_dir, summary_filename = create_output_directory_and_save_files(all_data, plan_name, 'single', files, enable_pdf, template_path)
        if not save_dir:
            return False
        run_chart_helper('education', plan_name, files, save_dir, summary_filename=summary_filename)
        return True
        
    except Exception as e:
        print_error(f"❌ 生成 {plan_name} 总结书失败: {e}")
        import traceback
        traceback.print_exc()
        return False

def parse_education_plan(text, usd_cny, idx, shared_data=None):
    """
    解析教育金方案数据 - 针对教育金模板的特定占位符格式
    
    Args:
        text: PDF提取的文本
        usd_cny: 美元兑人民币汇率
        idx: 方案索引 (0=单独模式, 1,2=对比模式)
        shared_data: 共享数据字典
        
    Returns:
        (data_dict, updated_shared_data): 解析的数据和更新的共享数据
    """
    if shared_data is None:
        shared_data = {}
    
    # 提取基本信息（只在第一次或单独模式时提取）
    if idx <= 1:
        # 姓名、年龄
        m = re.search(r"受保人姓名[:：]\s*([^\s]+)\s*(?:先生|女士).*?年龄[:：]\s*(\d+)", text)
        if m:
            shared_data["name"] = m.group(1).strip()
            shared_data["age"] = int(m.group(2))
            
        # 保险计划 & 供款期限
        m_plan = re.search(r"建[议議][书書]摘要[:：]\s*([^\(（]+)\s*[\(（]([0-9]+)\s*年[缴繳][费費]", text)
        if m_plan:
            shared_data["plan_name"] = m_plan.group(1).strip()
            shared_data["payment_term"] = m_plan.group(2).strip()
            
        # 计算衍生字段
        if "age" in shared_data:
            age = shared_data["age"]
            shared_data["age_plus_5"] = age + 5
            shared_data["age_plus_10"] = age + 10
            # 教育金特有的年限计算
            shared_data["years_withdraw_a"] = 20 - (age + 6)  # 20 - age_plus_5
            shared_data["years_withdraw_b"] = 102 - (age + 6)  # 102 - age_plus_5
    
    # 提取保费信息
    premium_usd = 0
    premium_total = _extract_total_premium(text)
    if premium_total:
        premium_usd = int(round(premium_total))
    else:
        m_row = None
        m_section = re.search(r"分红保单销售说明文件(.+?)(?=保障及利益摘要)", text, re.S)
        if m_section:
            part = m_section.group(1)
            m_row = re.search(r"首\s*12\s*个月意外身故赔偿[^\n]*", part)
        if m_row:
            row_text = m_row.group(0)
            nums = re.findall(r"\d{1,3}(?:,\d{3})*", row_text)
            if len(nums) >= 2:
                premium_usd = int(nums[1].replace(',', ''))
    premium_cny = round(premium_usd * usd_cny / 10000, 1)

    # 提取教育金现金提取数据 - 基于实际表格结构
    age_plus_5 = shared_data.get("age_plus_5", 7)
    
    # 从表格中提取age_plus_5年的提取金额
    withdraw_usd = 0
    # 查找现金提取表格中的 "age_plus_5 6 0 ... 数字" 格式的行
    pattern = fr"^\s*{age_plus_5}\s+6\s+0\s+[\d,]+\s+[\d,]+\s+([\d,]+)"
    m = re.search(pattern, text, re.M)
    if m:
        withdraw_usd = int(m.group(1).replace(",", ""))
    else:
        # 备用方案：查找任何包含age_plus_5和6的行，提取最后一个数字
        pattern = fr"^\s*{age_plus_5}\s+6\s+.*?([\d,]+)\s*$"
        m_backup = re.search(pattern, text, re.M)
        if m_backup:
            withdraw_usd = int(m_backup.group(1).replace(",", ""))
    
    withdraw_cny = withdraw_usd * usd_cny

    # 提取各年龄段的退保金额（从表格最后一列）
    cashout_data = {}
    target_ages = [18, 25, 45, 65, 85]
    
    for target_age in target_ages:
        # 查找 "年龄岁" 行的最后一个大数字（总退保金额）
        pattern = fr"^\s*{target_age}岁.*?([\d,]+)\s*$"
        m = re.search(pattern, text, re.M)
        if not m:
            # 尝试另一种格式
            pattern = fr"^\s*{target_age}\s+\d+.*?([\d,]+)\s+([\d,]+)\s*$"
            m = re.search(pattern, text, re.M)
            if m:
                # 取最后一个数字作为总额
                cashout_usd = int(m.group(2).replace(",", ""))
            else:
                cashout_usd = 0
        else:
            cashout_usd = int(m.group(1).replace(",", ""))
            
        cashout_data[target_age] = {
            'usd': cashout_usd,
            'cny': cashout_usd * usd_cny,
            'wan': cashout_usd * usd_cny / 10000
        }

    # 余额数据暂时使用退保金额（教育金模板可能不需要单独的余额）
    balance_data = {}
    balance_ages = [18, 35, 45, 55, 65]
    
    for balance_age in balance_ages:
        # 使用相近年龄的退保金额作为余额
        closest_age = min(target_ages, key=lambda x: abs(x - balance_age))
        balance_data[balance_age] = cashout_data[closest_age].copy()

    # 计算提取相关数据
    withdraw_cny_wan = round(withdraw_cny / 10000, 2) if withdraw_cny else 0
    withdraw_cny_month = withdraw_cny / 12 if withdraw_cny else 0
    
    years_withdraw_a = shared_data.get("years_withdraw_a", 0)
    years_withdraw_b = shared_data.get("years_withdraw_b", 0)
    
    # 计算总提取金额
    withdraw_cny_total_wan_a = withdraw_cny * years_withdraw_a / 10000 if withdraw_cny and years_withdraw_a > 0 else 0
    withdraw_cny_total_wan_b = withdraw_cny * years_withdraw_b / 10000 if withdraw_cny and years_withdraw_b > 0 else 0
    
    # 85岁金额的亿元计算
    cashout_cny_85_yi = round(cashout_data[85]['cny'] / 100000000) if cashout_data[85]['cny'] else 0

    # 生成数据字典 - 使用教育金模板的占位符格式
    data = {}
    
    # 基本信息
    data.update({
        "premium_usd": premium_usd,
        "premium_cny": premium_cny,
    })
    
    # 提取相关 - 教育金特有的a/b格式
    data.update({
        "withdraw_usd_a_1": withdraw_usd,
        "withdraw_usd_a_2": withdraw_usd,
        "withdraw_usd_a_3": withdraw_usd,
        "withdraw_usd_b": withdraw_usd,
        "withdraw_cny_wan_a_1": withdraw_cny_wan,
        "withdraw_cny_wan_a_2": withdraw_cny_wan,
        "withdraw_cny_wan_a_3": withdraw_cny_wan,
        "withdraw_cny_wan_b": withdraw_cny_wan,
        "withdraw_cny_month_a_1": withdraw_cny_month,
        "withdraw_cny_month_a_2": withdraw_cny_month,
        "withdraw_cny_month_a_3": withdraw_cny_month,
        "withdraw_cny_total_wan_a": withdraw_cny_total_wan_a,
        "withdraw_cny_total_wan_b": withdraw_cny_total_wan_b,
        "withdraw_cny_total_wan_a_1": withdraw_cny_total_wan_a,
        "withdraw_cny_total_wan_a_2": withdraw_cny_total_wan_a,
        "withdraw_cny_total_wan_a_3": withdraw_cny_total_wan_a,
        "years_withdraw_a": years_withdraw_a,
        "years_withdraw_b": years_withdraw_b,
    })
    
    # 退保金额 - 按具体年龄
    for age_key in [18, 25, 45, 65]:
        data[f"cashout_usd_{age_key}"] = cashout_data[age_key]['usd']
        data[f"cashout_cny_{age_key}_wan"] = cashout_data[age_key]['wan']
    
    # 85岁特殊处理
    data["cashout_cny_85_yi"] = cashout_cny_85_yi
    
    # 余额数据 - 按年龄和序号
    balance_mapping = {
        18: [(18, 2)],
        35: [(35, 1), (35, 2)],
        45: [(45, 2)],
        55: [(55, 1)],
        65: [(65, 2)]
    }
    
    for age_key, mappings in balance_mapping.items():
        for age_val, seq in mappings:
            data[f"balance_usd_{age_val}_{seq}"] = balance_data[age_key]['usd']
            data[f"balance_cny_{age_val}_wan_{seq}"] = balance_data[age_key]['wan']
    
    # 其他计算字段
    data["withdraw_multiple"] = round(cashout_data[85]['cny'] / (premium_usd * usd_cny), 1) if premium_usd else 0
    
    # 格式化数值（保留特定字段的小数）
    keep_decimal_fields = ["withdraw_cny_wan", "premium_cny", "cashout_cny_85_yi", "withdraw_multiple", "premium_usd"]
    for k, v in data.items():
        if isinstance(v, float) and not any(k.startswith(f) or f in k for f in keep_decimal_fields):
            try: 
                data[k] = int(v)
            except: 
                pass
    
    return data, shared_data

# 解析函数分派器
PARSE_FUNCTIONS = {
    'savings': parse_savings_plan,  # 使用新的储蓄险解析函数
    'critical_illness': parse_critical_illness_plan,
}

# ==============================================================================
# SECTION 4: 辅助函数 (汇率、Word生成、PDF转换等，基本保持不变)
# ==============================================================================
def get_usd_cny():
    print_info("正在获取美元兑人民币汇率...")
    apis = [
        ("https://v6.exchangerate-api.com/v6/EXCHANGE_RATE_API_KEY_REMOVED/pair/USD/CNY", lambda d: d.get("result") == "success" and d.get("conversion_rate")),
        ("http://api.exchangerate.host/live?access_key=EXCHANGERATE_HOST_API_KEY_PRIMARY_REMOVED&source=USD&currencies=CNY&format=1", lambda d: d.get("success") and d.get("quotes", {}).get("USDCNY")),
        ("http://api.exchangerate.host/live?access_key=EXCHANGERATE_HOST_API_KEY_SECONDARY_REMOVED&source=USD&currencies=CNY&format=1", lambda d: d.get("success") and d.get("quotes", {}).get("USDCNY"))
    ]
    for i, (url, extractor) in enumerate(apis, 1):
        try:
            print_info(f"  - 正在尝试 API {i}...")
            resp = requests.get(url, timeout=5)
            data = resp.json()
            rate = extractor(data)
            if rate:
                print_success(f"  - ✅ API {i} 获取成功: {float(rate):.4f}")
                return float(rate)
        except Exception as e:
            print_warn(f"  - ⚠️ API {i} 尝试失败: {e}")
    print_error("❌ 所有汇率API均获取失败，将使用默认值 7.25。")
    return 7.25

def extract_text(path):
    with pdfplumber.open(path) as pdf:
        text = "\n".join([p.extract_text() or "" for p in pdf.pages])
        return _decode_special_sequences(path, text, pdf)


def _format_value(value):
    if isinstance(value, int):
        return f"{value:,d}"
    if isinstance(value, float):
        dec = Decimal(str(value)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        text = format(dec, 'f')
        if '.' in text:
            text = text.rstrip('0').rstrip('.')
        # 添加千分位分隔符
        if '.' in text:
            int_part, dec_part = text.split('.')
            int_part = f"{int(int_part):,d}"
            text = f"{int_part}.{dec_part}"
        else:
            text = f"{int(text):,d}" if text else "0"
        return text
    if isinstance(value, Decimal):
        dec = value.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        text = format(dec, 'f')
        if '.' in text:
            text = text.rstrip('0').rstrip('.')
        # 添加千分位分隔符
        if '.' in text:
            int_part, dec_part = text.split('.')
            int_part = f"{int(int_part):,d}"
            text = f"{int_part}.{dec_part}"
        else:
            text = f"{int(text):,d}" if text else "0"
        return text
    return str(value)


def _get_decimal_value(data, keys):
    """从 data 中按顺序提取第一个合法的 Decimal 值"""
    for key in keys:
        if key not in data:
            continue
        raw = data.get(key)
        if raw in (None, '', 'null'):
            continue
        try:
            text = str(raw).replace(',', '')
            value = Decimal(text)
            return value
        except (InvalidOperation, ValueError):
            continue
    return None


def _format_premium_display(all_data):
    """生成“年交X万美金”中的 X 文本"""
    usd_keys = [
        'premium_usd', 'annual_premium_usd', 'premium_usd_0', 'premium_usd_1',
        'premium_usd_single', 'premium_usd_total', 'premium_usd_all'
    ]
    wan_keys = ['premium_usd_wan', 'annual_premium_usd_wan', 'premium_usd_10k']
    premium_usd = _get_decimal_value(all_data, usd_keys)
    if premium_usd is not None:
        premium_wan = premium_usd / Decimal('10000')
    else:
        premium_wan = _get_decimal_value(all_data, wan_keys)
    if premium_wan is None:
        return "0"
    premium_wan = premium_wan.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    text = format(premium_wan, 'f')
    if '.' in text:
        text = text.rstrip('0').rstrip('.')
    return text or "0"


def replace_text_in_paragraph(paragraph, data):
    """替换段落中的占位符，保持原有格式"""
    full_text = "".join(run.text for run in paragraph.runs)
    for key, value in data.items():
        placeholder = f"{{{key}}}"
        if placeholder in full_text:
            repl = _format_value(value)
            replaced = False
            # 优先尝试在单个run中替换
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, repl)
                    replaced = True
            # 如果占位符跨越多个runs，则合并处理
            if not replaced:
                combined = "".join(run.text for run in paragraph.runs)
                new_text = combined.replace(placeholder, repl)
                if combined != new_text:
                    for run in paragraph.runs:
                        run.text = ""
                    paragraph.runs[0].text = new_text

def generate_summary(template_path, output_path, data):
    if not Path(template_path).exists():
        print_error(f"❌ 模板文件不存在: {template_path}")
        return False
    doc = Document(template_path)
    for p in doc.paragraphs: replace_text_in_paragraph(p, data)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p, data)
    doc.save(output_path)
    print_success(f"✅ Word 已生成: {output_path}")
    return True

def create_output_directory_and_save_files(all_data, plan_name, mode, files, enable_pdf, template_path):
    """创建投保人姓名文件夹并保存文件（参考aia.py的方式）"""
    # 获取投保人姓名
    name = all_data.get('name', '未知')
    if not name or name == '未知':
        fallback = Path(files[0]).stem if files else '未命名'
        print_warn("⚠️ 无法获取投保人姓名，使用文件名作为文件夹名")
        name = fallback or '未命名'
    
    # 创建投保人姓名文件夹
    save_dir = Path.cwd() / name
    save_dir.mkdir(exist_ok=True)
    print_info(f"📁 创建输出文件夹: {save_dir}")
    
    def _pick_number(keys):
        for k in keys:
            v = all_data.get(k)
            if isinstance(v, (int, float)):
                return float(v)
            try:
                return float(str(v))
            except Exception:
                continue
        return None

    age_val = _pick_number(['age', 'insured_age', 'applicant_age'])
    age_str = f"{int(age_val)}岁" if isinstance(age_val, (int, float)) and age_val is not None else '年龄未知'

    if plan_name == '储蓄险':
        # 对比模式时，显示两个方案的保费
        if mode == 'comparison':
            premium_1 = all_data.get('premium_usd_1', 0)
            premium_2 = all_data.get('premium_usd_2', 0)
            # 转换为万元显示
            premium_wan_1 = premium_1 / 10000 if premium_1 else 0
            premium_wan_2 = premium_2 / 10000 if premium_2 else 0
            # 格式化保费显示（去掉多余的0）
            def _fmt_wan(v):
                text = f"{v:.2f}".rstrip('0').rstrip('.')
                return text or "0"
            premium_text = f"{_fmt_wan(premium_wan_1)}vs{_fmt_wan(premium_wan_2)}"
        else:
            premium_text = _format_premium_display(all_data)
        x_val = f"年交{premium_text}万美金"
        plan_label = '财富增值方案'
    else:
        coverage_wan = _pick_number(['coverage_usd_wan', 'sum_assured_usd_wan', 'coverage_usd_10k', 'coverage_usd', 'sum_assured_usd'])
        if coverage_wan is None:
            coverage_cny = _pick_number(['coverage_cny_wan', 'sum_assured_cny_wan', 'coverage_cny'])
            if coverage_cny is not None:
                usd_cny = float(os.environ.get('USD_CNY_RATE', '7.2'))
                coverage_wan = coverage_cny / usd_cny
        if coverage_wan is not None and coverage_wan > 1000:
            coverage_wan = coverage_wan / 10000.0
        x_val = f"{(coverage_wan if coverage_wan is not None else 0):g}万美金保额"
        plan_label = '重大疾病保障方案'

    # 根据模式生成不同的文件名：单独总结书 vs 对比总结书
    summary_type = "对比总结书" if mode == 'comparison' else "总结书"
    output_name = f"【{plan_name}{summary_type}】{name}_{age_str}_{plan_label}（{x_val}）.docx"
    
    # 生成Word文档路径
    final_docx_path = save_dir / output_name
    
    # 生成Word文档
    try:
        success = generate_summary(template_path, final_docx_path, all_data)
        if not success:
            print_error(f"❌ 生成 {plan_name} 总结书失败: 模板处理失败")
            return None, None
        
        # PDF转换
        if enable_pdf:
            print_info("正在转换为PDF...")
            pdf_path = convert_to_pdf(str(final_docx_path))
            if pdf_path:
                # 将PDF移动到投保人文件夹
                final_pdf_name = output_name.replace(".docx", ".pdf")
                final_pdf_path = save_dir / final_pdf_name
                import shutil
                shutil.move(pdf_path, final_pdf_path)
                print_success(f"✅ {plan_name} 总结书已完成: {final_docx_path.name} & {final_pdf_path.name}")
            else:
                print_success(f"✅ {plan_name} 总结书已完成: {final_docx_path.name}")
                print_info("💡 您可以手动打开Word文档并导出为PDF")
        else:
            print_success(f"✅ {plan_name} 总结书已完成: {final_docx_path.name}")
        
        return save_dir, final_docx_path.name
        
    except Exception as e:
        print_error(f"❌ 生成 {plan_name} 总结书失败: {e}")
        return None, None

# PDF转换函数 (convert_docx_to_pdf_... 等) 保持原样，此处省略以节约篇幅...
# ... (此处应粘贴原始脚本中的所有PDF转换函数)
def convert_docx_to_pdf_using_pages(docx_path, pdf_path):
    """使用 Pages 转换 PDF（参考 aia.py 的简单实现）"""
    apple_script = f'''
    tell application "Pages"
        open POSIX file "{docx_path}"
        delay 2
        tell front document
            export to POSIX file "{pdf_path}" as PDF
            close saving no
        end tell
    end tell
    '''
    try:
        print_info("使用 Pages 转换 PDF...")
        subprocess.run(["osascript", "-e", apple_script], check=True)
        return Path(pdf_path).exists()
    except Exception as e:
        print_warn(f"Pages 转换失败: {e}")
    return False

def convert_docx_to_pdf_using_docx2pdf(docx_path, pdf_path):
    """使用 docx2pdf 转换 PDF（参考 aia.py 的简单实现）"""
    if SYSTEM == "Darwin":
        try:
            res = subprocess.run(
                ["mdfind", "kMDItemCFBundleIdentifier == 'com.microsoft.Word'"],
                capture_output=True, text=True)
            if not res.stdout.strip():
                print_warn("未检测到 Microsoft Word，跳过 docx2pdf。")
                return False
        except Exception:
            return False
    try:
        print_info("使用 docx2pdf 转换 PDF...")
        convert(docx_path, pdf_path)
        return Path(pdf_path).exists()
    except Exception as e:
        print_warn(f"docx2pdf 转换失败: {e}")
    return False

def convert_docx_to_pdf_using_libreoffice(docx_path, pdf_path):
    """使用 LibreOffice 转换 PDF（参考 aia.py 的简单实现）"""
    soffice_path = shutil.which("soffice")
    if not soffice_path:
        print_warn("未检测到 LibreOffice")
        return False
    try:
        print_info("使用 LibreOffice 转换 PDF...")
        subprocess.run([soffice_path, "--headless", "--convert-to", "pdf",
                        "--outdir", str(Path(pdf_path).parent), str(docx_path)], check=True)
        return Path(pdf_path).exists()
    except Exception as e:
        print_warn(f"LibreOffice 转换失败: {e}")
    return False

def convert_docx_to_pdf_using_textutil(docx_path, pdf_path):
    """使用 textutil 转换 PDF（低保真）"""
    try:
        print_info("使用 textutil 转换 PDF(低保真)...")
        tmp_rtf = Path(str(docx_path).replace(".docx", ".rtf"))
        subprocess.run(["textutil", "-convert", "rtf", str(docx_path)], check=True)
        subprocess.run(["textutil", "-convert", "pdf", str(tmp_rtf), "-output", str(pdf_path)], check=True)
        if tmp_rtf.exists(): tmp_rtf.unlink()
        return Path(pdf_path).exists()
    except Exception as e:
        print_warn(f"textutil 转换失败: {e}")
    return False

def convert_docx_to_pdf_mac_priority(docx_path, pdf_path):
    """macOS 优先的 PDF 转换策略（优化版）"""
    if SYSTEM == "Darwin":
        # 优先使用LibreOffice（更稳定）
        if convert_docx_to_pdf_using_libreoffice(docx_path, pdf_path): return True
        if convert_docx_to_pdf_using_pages(docx_path, pdf_path): return True
        if convert_docx_to_pdf_using_docx2pdf(docx_path, pdf_path): return True
        return convert_docx_to_pdf_using_textutil(docx_path, pdf_path)
    else:
        if convert_docx_to_pdf_using_libreoffice(docx_path, pdf_path): return True
        return convert_docx_to_pdf_using_docx2pdf(docx_path, pdf_path)

def convert_to_pdf(docx_path):
    """尝试将 Word 文档转换为 PDF，使用多种方法确保成功率"""
    pdf_path = str(docx_path).replace('.docx', '.pdf')
    
    print_info("开始PDF转换...")
    
    # 使用与 aia.py 相同的转换策略
    if convert_docx_to_pdf_mac_priority(docx_path, pdf_path):
        return pdf_path
    
    print_warn(f"⚠️ 所有PDF转换方法均失败，仅生成了 Word 文档: {docx_path}")
    print_info("💡 建议：")
    print_info("  1. 确保已安装 Pages 或 Microsoft Word")
    print_info("  2. 检查文件访问权限设置")
    print_info("  3. 或手动安装 LibreOffice 作为备用")
    return None

# ==============================================================================
# SECTION 5: 制图小助手 (新增)
# ==============================================================================
def _find_detail_page_index(pdf_path):
    """查找包含“详细说明”关键字的页面，找不到则回退到 P.12"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            if not pdf.pages:
                return None
            fallback_idx = min(11, len(pdf.pages) - 1)
            for idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if any(keyword in text for keyword in DETAIL_SECTION_KEYWORDS):
                    return idx
            return fallback_idx
    except Exception as e:
        print_warn(f"制图小助手无法读取 {Path(pdf_path).name}: {e}")
        return None


def _get_overlay_dimensions():
    """缓存并返回叠加图宽高"""
    global _OVERLAY_DIMENSIONS
    if _OVERLAY_DIMENSIONS is not None:
        return _OVERLAY_DIMENSIONS
    if not ANNOTATION_OVERLAY_PATH.exists():
        return None
    pix = None
    try:
        pix = fitz.Pixmap(str(ANNOTATION_OVERLAY_PATH))
        _OVERLAY_DIMENSIONS = (pix.width, pix.height)
        return _OVERLAY_DIMENSIONS
    except Exception as e:
        print_warn(f"制图小助手无法读取叠加模板: {e}")
        return None
    finally:
        if pix:
            pix = None


def _calculate_overlay_rect(page_rect, overlay_size, cfg):
    """根据配置计算叠加图的摆放矩形"""
    overlay_w, overlay_h = overlay_size
    overlay_ratio = overlay_w / overlay_h if overlay_h else 1
    fit_mode = (cfg.get('fit') or OVERLAY_SETTINGS['default']['fit']).lower()

    def contain():
        width = page_rect.width
        height = width / overlay_ratio if overlay_ratio else page_rect.height
        if height > page_rect.height:
            height = page_rect.height
            width = height * overlay_ratio if overlay_ratio else page_rect.width
        return width, height

    def cover():
        width = page_rect.width
        height = width / overlay_ratio if overlay_ratio else page_rect.height
        if height < page_rect.height:
            height = page_rect.height
            width = height * overlay_ratio if overlay_ratio else page_rect.width
        return width, height

    if fit_mode == 'width':
        target_w = page_rect.width
        target_h = target_w / overlay_ratio if overlay_ratio else page_rect.height
    elif fit_mode == 'height':
        target_h = page_rect.height
        target_w = target_h * overlay_ratio if overlay_ratio else page_rect.width
    elif fit_mode == 'cover':
        target_w, target_h = cover()
    else:  # contain/default
        target_w, target_h = contain()

    anchor = (cfg.get('anchor') or OVERLAY_SETTINGS['default']['anchor']).lower()
    x = page_rect.x0
    y = page_rect.y0

    horizontal = 'center'
    vertical = 'center'
    if '-' in anchor:
        parts = anchor.split('-', 1)
        vertical, horizontal = parts[0], parts[1]
    else:
        if anchor in ('top', 'bottom', 'center', 'middle'):
            vertical = anchor if anchor in ('top', 'bottom') else 'center'
        if anchor in ('left', 'right', 'center', 'middle'):
            horizontal = anchor if anchor in ('left', 'right') else 'center'

    if horizontal == 'left':
        x = page_rect.x0
    elif horizontal == 'right':
        x = page_rect.x1 - target_w
    else:
        x = page_rect.x0 + (page_rect.width - target_w) / 2

    if vertical == 'top':
        y = page_rect.y0
    elif vertical == 'bottom':
        y = page_rect.y1 - target_h
    else:
        y = page_rect.y0 + (page_rect.height - target_h) / 2

    x += cfg.get('offset_x', 0)
    y += cfg.get('offset_y', 0)
    return fitz.Rect(x, y, x + target_w, y + target_h)


def _export_page_with_overlay(source_pdf, page_index, output_pdf_path, overlay_cfg):
    """导出单页 PDF 并叠加标注图层"""
    doc = helper_doc = None
    try:
        doc = fitz.open(source_pdf)
        if doc.page_count == 0:
            print_warn(f"制图小助手无法处理空白 PDF: {source_pdf}")
            return False
        if page_index is None:
            page_index = min(11, doc.page_count - 1)
        page_index = max(0, min(page_index, doc.page_count - 1))
        helper_doc = fitz.open()
        helper_doc.insert_pdf(doc, from_page=page_index, to_page=page_index)
        page = helper_doc[0]
        overlay_size = _get_overlay_dimensions()
        if not overlay_size:
            return False
        target_rect = _calculate_overlay_rect(page.rect, overlay_size, overlay_cfg)
        page.insert_image(target_rect, filename=str(ANNOTATION_OVERLAY_PATH),
                          overlay=True, keep_proportion=True)
        helper_doc.save(output_pdf_path)
        print_success(f"✅ 制图小助手输出: {output_pdf_path}")
        return True
    except Exception as e:
        print_warn(f"制图小助手叠加失败（{Path(source_pdf).name}）: {e}")
        return False
    finally:
        if helper_doc:
            helper_doc.close()
        if doc:
            doc.close()


def run_chart_helper(plan_type, plan_name, files, save_dir, summary_filename=None):
    """在储蓄险流程末尾触发制图小助手"""
    if plan_type != 'savings':
        return
    if not ANNOTATION_OVERLAY_PATH.exists():
        print_warn("⚠️ 制图小助手模板未找到，已跳过叠加步骤")
        return
    
    # 默认运行制图小助手，不再询问用户
    print_info(f"🎨 正在运行制图小助手，为{plan_name}生成带标注的详细说明单页...")

    generated = False
    overlay_cfg = OVERLAY_SETTINGS.get(plan_type, OVERLAY_SETTINGS['default'])
    for pdf_path in files:
        if not Path(pdf_path).exists():
            continue
        page_index = _find_detail_page_index(pdf_path)
        if page_index is None:
            continue
        if summary_filename:
            summary_stem = Path(summary_filename).stem
            output_pdf_path = save_dir / f"{summary_stem}_投资总览图.pdf"
        else:
            output_pdf_path = save_dir / f"{Path(pdf_path).stem}_投资总览图.pdf"
        if _export_page_with_overlay(pdf_path, page_index, output_pdf_path, overlay_cfg):
            generated = True

    if not generated:
        print_warn("⚠️ 制图小助手未找到可用页面，未生成标注单页")

# ==============================================================================
# SECTION 6: 主流程执行器 (新增)
# ==============================================================================
def execute_single_task(task, usd_cny, enable_pdf=False):
    """执行单个任务（单独总结书或对比总结书）"""
    plan_type = task['type']
    mode = task['mode']
    files = task['files']
    
    config = PLAN_CONFIG[plan_type]
    plan_name = config['name']
    template_path = config['templates'].get(mode)
    parse_func = PARSE_FUNCTIONS[plan_type]

    # 智能模板回退机制
    if not template_path:
        if mode == 'comparison' and config['templates'].get('single'):
            fallback_template = config['templates']['single']
            print_warn(f"⚠️ 未配置对比模板，自动使用单独模板: {fallback_template}")
            template_path = fallback_template
            mode = 'single'
        else:
            print_error("❌ 未找到可用模板配置")
            return False

    template_exists = Path(template_path).exists()
    if not template_exists:
        if mode == 'comparison' and config['templates'].get('single'):
            fallback_template = config['templates']['single']
            if Path(fallback_template).exists():
                print_warn(f"⚠️ 对比模板不存在，自动使用单独模板: {fallback_template}")
                template_path = fallback_template
                mode = 'single'  # 更新模式为单独处理
            else:
                print_error(f"❌ 模板文件不存在: {template_path}")
                print_error(f"❌ 回退模板也不存在: {fallback_template}")
                return False
        else:
            print_error(f"❌ 模板文件不存在: {template_path}")
            return False
    
    print_info(f"🔄 正在处理 {plan_name} - {'对比总结书' if mode == 'comparison' else '单独总结书'}")
    print(f"   文件: {', '.join([Path(f).name for f in files])}")
    print(f"   模板: {template_path}")
    
    # 常规处理其他保险类型
    # 提取数据
    all_data = {}
    shared_data = {}
    
    
    for idx, file_path in enumerate(files, 1):
        # 验证文件是否存在（防止任务确定后文件被移动或删除）
        if not Path(file_path).exists():
            print_error(f"    ❌ 文件不存在: {file_path}")
            print_warn(f"    提示：任务创建后文件可能被移动或删除")
            return False
            
        meta = FILE_METADATA.get(str(Path(file_path).resolve()))
        if meta:
            if meta.get('name') and 'name' not in shared_data:
                shared_data['name'] = meta['name']
            if meta.get('age') is not None and 'age' not in shared_data:
                shared_data['age'] = meta['age']
            if meta.get('payment_term') is not None and 'payment_term' not in shared_data:
                shared_data['payment_term'] = meta['payment_term']
            if meta.get('plan_name') and 'plan_name' not in shared_data:
                shared_data['plan_name'] = meta['plan_name']
        
        print_info(f"  - 正在提取文件 {idx}: {Path(file_path).name}")
        try:
            text = extract_text(file_path)
            if mode == 'single':
                # 单独方案：每个文件都用 idx=0（不加数字后缀）
                data, shared_data = parse_func(text, usd_cny, 0, shared_data)
            else:
                # 对比方案：使用文件索引作为后缀 (1, 2)
                data, shared_data = parse_func(text, usd_cny, idx, shared_data)
            all_data.update(data)
        except Exception as e:
            print_error(f"    ❌ 提取文件 {Path(file_path).name} 失败: {e}")
            return False
    
    # 合并共享数据
    all_data.update(shared_data)
    
    # 储蓄险年龄>= 45岁时，使用特殊模板
    if plan_type == 'savings' and mode == 'single':
        customer_age = shared_data.get('age')
        if customer_age is not None and customer_age >= 45:
            age_template = 'templates/template_savings_standalone_45.docx'
            if Path(age_template).exists():
                print_info(f"📋 检测到客户年龄 {customer_age} 岁 (≥45)，使用45岁专用模板")
                template_path = age_template
            else:
                print_warn(f"⚠️ 45岁专用模板不存在: {age_template}，使用默认模板")
    
    # 使用新的文件保存方式（创建投保人姓名文件夹）
    save_dir, summary_filename = create_output_directory_and_save_files(all_data, plan_name, mode, files, enable_pdf, template_path)
    if not save_dir:
        return False
    # 只有单独总结书才生成投资总览图，对比总结书不生成
    if mode == 'single':
        run_chart_helper(plan_type, plan_name, files, save_dir, summary_filename=summary_filename)
    return True

def check_pdf_conversion_capabilities():
    """检查PDF转换能力并给出建议"""
    capabilities = []
    suggestions = []
    
    # 检查 Pages
    if SYSTEM == "Darwin":
        try:
            result = subprocess.run(["mdfind", "kMDItemCFBundleIdentifier == 'com.apple.iWork.Pages'"], 
                                  capture_output=True, text=True, timeout=5)
            if result.stdout.strip():
                capabilities.append("Pages ✅")
            else:
                suggestions.append("安装 Pages（推荐，最稳定）")
        except:
            pass
    
    # 检查 Microsoft Word
    if SYSTEM == "Darwin":
        try:
            result = subprocess.run(["mdfind", "kMDItemCFBundleIdentifier == 'com.microsoft.Word'"], 
                                  capture_output=True, text=True, timeout=5)
            if result.stdout.strip():
                capabilities.append("Microsoft Word ✅")
            else:
                suggestions.append("安装 Microsoft Word")
        except:
            pass
    
    # 检查 LibreOffice
    soffice_path = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice_path and SYSTEM == "Darwin":
        if Path("/Applications/LibreOffice.app/Contents/MacOS/soffice").exists():
            soffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    
    if soffice_path:
        capabilities.append("LibreOffice ✅")
    else:
        suggestions.append("安装 LibreOffice（免费备选方案）")
    
    return capabilities, suggestions

def execute_all_tasks(tasks):
    """执行所有任务"""
    if not tasks:
        print_warn("⚠️ 没有要执行的任务")
        return
    
    print_info(f"🚀 开始执行 {len(tasks)} 个任务...")
    
    # 检查PDF转换能力
    capabilities, suggestions = check_pdf_conversion_capabilities()
    
    if capabilities:
        print_info(f"📄 PDF转换能力: {', '.join(capabilities)}")
    else:
        print_warn("⚠️ 未检测到PDF转换软件")
        if suggestions:
            print_info("💡 建议安装以下软件之一以支持PDF转换:")
            for suggestion in suggestions:
                print_info(f"   • {suggestion}")
    
    # 默认启用PDF转换（如果有转换能力的话）
    if capabilities:
        print()
        print_info("📄 已启用自动PDF转换")
        enable_pdf = True
    else:
        enable_pdf = False
        print_info("将仅生成Word文档")
    
    # 获取汇率（只需获取一次）
    usd_cny = get_usd_cny()
    
    success_count = 0
    total_count = len(tasks)
    
    for i, task in enumerate(tasks, 1):
        print_info(f"\n📋 任务 {i}/{total_count}")
        print("─" * 50)
        
        if execute_single_task(task, usd_cny, enable_pdf):
            success_count += 1
        else:
            print_error(f"❌ 任务 {i} 执行失败")
    
    # 总结执行结果
    print("\n" + "═" * 60)
    if success_count == total_count:
        print_success(f"🎉 所有任务执行完成！成功 {success_count}/{total_count}")
    else:
        print_warn(f"⚠️ 部分任务执行完成。成功 {success_count}/{total_count}")
    print("═" * 60)

def main():
    """主函数"""
    print_info("🏛️ AIA 保险方案总结书生成器")
    print_info("支持储蓄险、重疾险方案的自动化处理\n")
    
    # 检查PDF转换工具
    print_info("=" * 60)
    check_pdf_conversion_tools()
    print_info("=" * 60)
    print()
    
    try:
        # 扫描和分类 PDF 文件
        classified_pdfs, file_metadata = scan_and_classify_pdfs()
        
        # 确定要执行的任务
        tasks = determine_tasks(classified_pdfs, file_metadata)