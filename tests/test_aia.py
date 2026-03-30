import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from fastapi.testclient import TestClient
from docx import Document


os.environ.setdefault("AIA_SKIP_VENV", "1")
os.environ.setdefault("AIA_AUTO_INSTALL", "0")

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
os.environ.setdefault("BIBHELPER_DATA_ROOT", str(ROOT / ".test-service-data"))

import aia
from apps.service.app.main import ServiceConfig, create_app
from bib_core import GeneratedArtifact, RunResult
from bib_core import core as core_module
import bib_core.env_loader as env_loader_module


class FakePage:
    def __init__(self, text):
        self._text = text

    def extract_text(self):
        return self._text


class FakePdf:
    def __init__(self, pages):
        self.pages = [FakePage(text) for text in pages]

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False


class FakeResponse:
    def __init__(self, payload, status_code=200):
        self._payload = payload
        self._status_code = status_code

    def raise_for_status(self):
        if self._status_code >= 400:
            raise RuntimeError(f"http {self._status_code}")

    def json(self):
        return self._payload


class NameExtractionTests(unittest.TestCase):
    def test_extract_name_and_age_allows_spaces(self):
        text = "受保人姓名: Mary Jane 女士 年龄: 31"
        self.assertEqual(aia._extract_name_and_age_from_text(text), ("Mary Jane", 31))

    def test_extract_name_and_age_does_not_capture_label_text(self):
        text = "受保人姓名: Mary Jane\n年龄: 31"
        self.assertEqual(aia._extract_name_and_age_from_text(text), ("Mary Jane", 31))

    def test_extract_payment_term_and_age_preserves_space_name(self):
        fake_pdf = FakePdf([
            "受保人姓名: Mary Jane 女士 年龄: 31\n5年缴费"
        ])
        with mock.patch.object(aia.pdfplumber, "open", return_value=fake_pdf):
            with mock.patch.object(aia, "_decode_special_sequences", side_effect=lambda path, text, pdf=None: text):
                payment_term, age, name = aia.extract_payment_term_and_age("fake.pdf")

        self.assertEqual(payment_term, 5)
        self.assertEqual(age, 31)
        self.assertEqual(name, "Mary Jane")


class SavingsTaskTests(unittest.TestCase):
    def test_build_savings_tasks_creates_comparison_for_same_person(self):
        files = [
            str(ROOT / "sample_a.pdf"),
            str(ROOT / "sample_b.pdf"),
        ]
        metadata = {
            files[0]: {"name": "Mary Jane", "age": 31},
            files[1]: {"name": "Mary Jane", "age": 31},
        }

        tasks = aia._build_savings_tasks(files, metadata)

        self.assertIn(
            {
                "type": "savings",
                "mode": "comparison",
                "files": sorted(files),
            },
            tasks,
        )


class ScopeRegressionTests(unittest.TestCase):
    def test_current_scope_excludes_education(self):
        self.assertNotIn("education", aia.PLAN_CONFIG)
        self.assertEqual(set(aia.PARSE_FUNCTIONS), {"savings", "critical_illness"})
        self.assertEqual(aia.classify_by_payment_term_and_age(5, 10, "儿童方案.pdf"), "savings")
        self.assertIsNone(aia.classify_by_payment_term_and_age(None, 10, "教育金方案.pdf"))


class HtmlRenderingTests(unittest.TestCase):
    def test_create_output_uses_docx_pdf_pipeline(self):
        with tempfile.TemporaryDirectory() as tempdir:
            temp_root = Path(tempdir)
            template_docx = temp_root / "template_savings_standalone.docx"
            source_pdf = temp_root / "plan.pdf"
            source_pdf.write_bytes(b"%PDF-1.4\n")

            document = Document()
            document.add_paragraph("客户：{name}")
            document.add_paragraph("年龄：{age}")
            document.save(template_docx)

            def fake_generate_summary(template_path, output_path, data):
                generated = Document()
                generated.add_paragraph(f"客户：{data['name']}")
                generated.save(output_path)
                return True

            generated_pdf = temp_root / "generated.pdf"
            generated_pdf.write_bytes(b"%PDF-1.4\ndocx-pdf")

            with mock.patch.object(core_module, "generate_summary", side_effect=fake_generate_summary):
                with mock.patch.object(core_module, "convert_html_to_pdf_using_weasyprint") as html_pdf_mock:
                    with mock.patch.object(core_module, "convert_to_pdf") as docx_pdf_mock:
                        docx_pdf_mock.return_value = str(generated_pdf)
                        save_dir, summary_name, generated_paths = core_module.create_output_directory_and_save_files(
                            {"name": "Mary Jane", "age": 31},
                            "储蓄险",
                            "single",
                            [str(source_pdf)],
                            True,
                            str(template_docx),
                            output_root=temp_root / "output",
                        )

            self.assertIsNotNone(save_dir)
            self.assertTrue(summary_name.endswith(".docx"))
            suffixes = {Path(path).suffix for path in generated_paths}
            self.assertEqual(suffixes, {".docx", ".pdf"})
            html_pdf_mock.assert_not_called()
            docx_pdf_mock.assert_called_once()


class PremiumExtractionTests(unittest.TestCase):
    SAMPLE_SUMMARY_TEXT = """
受保人姓名: Mary Jane 女士 年龄: 45
建议书摘要: 财富增值方案 (25年缴费)
投保时年缴总保费：5,054
保险业监管局(IA)保费征费：0
分红保单销售说明文件
3. 基本计划 - 说明摘要
保单年度终结 缴付保费总额 退保发还金额 严重疾病赔偿 / 身故赔偿额
1 5,054 0 0 0 225,000 0 225,000
5 25,268 837 60 897 225,000 90 225,090
10 50,535 5,423 180 5,603 225,000 270 225,270
15 75,803 16,890 300 17,190 150,000 420 150,420
20 101,070 38,282 540 38,822 150,000 720 150,720
25 126,338 51,807 6,930 58,737 150,000 8,820 158,820
30 126,338 57,176 106,050 163,226 150,000 129,375 279,375
    65岁 126,338 65,318 147,945 213,263 150,000 173,385 323,385
以上摘要说明：请参考说明部分。
""".strip()

    def test_extract_policy_total_premium_from_summary_table(self):
        total = core_module._extract_policy_total_premium(self.SAMPLE_SUMMARY_TEXT)
        self.assertEqual(total, 126338)

    def test_parse_savings_plan_uses_extracted_total_premium(self):
        data, shared_data = core_module.parse_savings_plan(self.SAMPLE_SUMMARY_TEXT, 7.2, 0, {})
        self.assertEqual(shared_data["payment_term"], "25")
        self.assertEqual(data["premium_usd_0"], 5054)
        self.assertEqual(data["premium_usd_all"], 126338)
        self.assertNotEqual(data["premium_usd_all"], 126350)
        self.assertEqual(data["premium_cny_all_wan"], 91.0)

    def test_parse_critical_illness_plan_uses_extracted_total_premium(self):
        ci_text = self.SAMPLE_SUMMARY_TEXT.replace("财富增值方案", "爱伴航")
        data, shared_data = core_module.parse_critical_illness_plan(ci_text, 7.2, 0, {})
        self.assertEqual(shared_data["payment_term"], "25")
        self.assertEqual(data["premium_usd_0"], 5054)
        self.assertEqual(data["premium_usd_all"], 126338)
        self.assertNotEqual(data["premium_usd_all"], 126350)
        self.assertEqual(data["premium_cny_all_wan"], 91.0)


class EnvLoaderTests(unittest.TestCase):
    def test_load_repo_env_prefers_process_env_then_deploy_env(self):
        with tempfile.TemporaryDirectory() as tempdir:
            repo_root = Path(tempdir)
            deploy_dir = repo_root / "deploy"
            deploy_dir.mkdir(parents=True, exist_ok=True)
            (deploy_dir / ".env.runtime").write_text(
                "USD_CNY_RATE=6.8\nEXCHANGERATE_HOST_API_KEY_PRIMARY=runtime-key\n",
                encoding="utf-8",
            )
            (deploy_dir / ".env").write_text(
                "USD_CNY_RATE=6.7\nEXCHANGE_RATE_API_KEY=env-key\n",
                encoding="utf-8",
            )

            with mock.patch.dict(os.environ, {"USD_CNY_RATE": "6.9"}, clear=True):
                loaded = env_loader_module.load_repo_env(repo_root)
                self.assertEqual(
                    [path.resolve() for path in loaded],
                    [(deploy_dir / ".env.runtime").resolve(), (deploy_dir / ".env").resolve()],
                )
                self.assertEqual(os.environ["USD_CNY_RATE"], "6.9")
                self.assertEqual(os.environ["EXCHANGERATE_HOST_API_KEY_PRIMARY"], "runtime-key")
                self.assertEqual(os.environ["EXCHANGE_RATE_API_KEY"], "env-key")


class ExchangeRateTests(unittest.TestCase):
    def test_get_usd_cny_uses_public_channel_when_private_keys_missing(self):
        expected_url = "https://api.frankfurter.dev/v2/rate/USD/CNY"

        def fake_get(url, timeout):
            self.assertEqual(url, expected_url)
            self.assertEqual(timeout, 5)
            return FakeResponse({"rate": 6.88})

        with mock.patch.dict(os.environ, {"USD_CNY_RATE": "6.9"}, clear=True):
            with mock.patch.object(core_module.requests, "get", side_effect=fake_get) as get_mock:
                rate = core_module.get_usd_cny()

        self.assertEqual(rate, 6.88)
        self.assertEqual(get_mock.call_count, 1)

    def test_get_usd_cny_falls_back_to_public_channel_after_private_api_failure(self):
        urls = []

        def fake_get(url, timeout):
            urls.append(url)
            self.assertEqual(timeout, 5)
            if "v6.exchangerate-api.com" in url:
                raise RuntimeError("boom")
            return FakeResponse({"rate": 6.91})

        with mock.patch.dict(
            os.environ,
            {
                "EXCHANGE_RATE_API_KEY": "test-key",
                "USD_CNY_RATE": "6.9",
            },
            clear=True,
        ):
            with mock.patch.object(core_module.requests, "get", side_effect=fake_get):
                rate = core_module.get_usd_cny()

        self.assertEqual(rate, 6.91)
        self.assertEqual(
            urls,
            [
                "https://v6.exchangerate-api.com/v6/test-key/pair/USD/CNY",
                "https://api.frankfurter.dev/v2/rate/USD/CNY",
            ],
        )

    def test_get_usd_cny_uses_6_point_9_default_when_all_channels_fail(self):
        with mock.patch.dict(os.environ, {"USD_CNY_RATE": "6.9"}, clear=True):
            with mock.patch.object(core_module.requests, "get", side_effect=RuntimeError("timeout")):
                rate = core_module.get_usd_cny()

        self.assertEqual(rate, 6.9)


class ServiceTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.data_root = Path(self.tempdir.name)
        self.config = ServiceConfig(
            data_root=self.data_root,
            job_retention_days=7,
            max_upload_files=5,
            max_upload_bytes=50 * 1024 * 1024,
            max_concurrent_jobs=1,
            shortcut_api_token="token-123",
            web_admin_password="pass-123",
            session_secret="secret-123",
            templates_dir=ROOT / "apps/service/templates",
            static_dir=ROOT / "apps/service/static",
            enable_pdf=False,
        )
        self.app = create_app(config=self.config, processor=self.fake_processor)
        self.client = TestClient(self.app)

    def tearDown(self):
        self.client.close()
        self.tempdir.cleanup()

    @staticmethod
    def fake_pdf_upload(name="sample.pdf"):
        return (name, b"%PDF-1.4\n%fake pdf\n", "application/pdf")

    @staticmethod
    def fake_processor(options):
        customer_dir = options.output_root / "Mary Jane"
        customer_dir.mkdir(parents=True, exist_ok=True)
        docx_path = customer_dir / "summary.docx"
        pdf_path = customer_dir / "summary.pdf"
        docx_path.write_text("ok", encoding="utf-8")
        pdf_path.write_bytes(b"%PDF-1.4\nartifact")
        return RunResult(
            job_id=None,
            classified={"savings": [path.name for path in options.input_files]},
            tasks=[{"type": "savings", "mode": "single", "files": [str(options.input_files[0])]}],
            artifacts=[
                GeneratedArtifact(
                    relative_path=str(docx_path.relative_to(options.output_root)),
                    kind="docx",
                    customer_name="Mary Jane",
                    plan_type="savings",
                    source_filenames=[path.name for path in options.input_files],
                ),
                GeneratedArtifact(
                    relative_path=str(pdf_path.relative_to(options.output_root)),
                    kind="pdf",
                    customer_name="Mary Jane",
                    plan_type="savings",
                    source_filenames=[path.name for path in options.input_files],
                ),
            ],
            warnings=[],
        )

    def login(self):
        response = self.client.post("/login", data={"password": "pass-123"}, follow_redirects=False)
        self.assertEqual(response.status_code, 303)

    def test_api_requires_bearer_token(self):
        response = self.client.post(
            "/api/v1/process",
            files=[("files", self.fake_pdf_upload())],
        )
        self.assertEqual(response.status_code, 401)

    def test_api_process_returns_zip_and_job_metadata(self):
        response = self.client.post(
            "/api/v1/process",
            headers={"Authorization": "Bearer token-123"},
            files=[("files", self.fake_pdf_upload("proposal.pdf"))],
        )
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["content-type"], "application/zip")
        job_id = response.headers["x-job-id"]

        status_response = self.client.get(
            f"/api/v1/jobs/{job_id}",
            headers={"Authorization": "Bearer token-123"},
        )
        self.assertEqual(status_response.status_code, 200)
        self.assertEqual(status_response.json()["status"], "completed")
        self.assertEqual(status_response.json()["result_filename"], "Mary Jane保险总结书.zip")
        self.assertTrue(status_response.json()["created_at"].endswith("+08:00"))
        self.assertTrue(status_response.json()["expires_at"].endswith("+08:00"))

        download_response = self.client.get(
            f"/api/v1/jobs/{job_id}/download",
            headers={"Authorization": "Bearer token-123"},
        )
        self.assertEqual(download_response.status_code, 200)
        self.assertEqual(download_response.headers["content-type"], "application/zip")

    def test_web_requires_login_then_lists_history(self):
        response = self.client.get("/jobs", follow_redirects=False)
        self.assertEqual(response.status_code, 303)
        self.assertEqual(response.headers["location"], "/login")

        self.login()
        upload_response = self.client.post(
            "/upload",
            files=[("files", self.fake_pdf_upload("history.pdf"))],
            follow_redirects=False,
        )
        self.assertEqual(upload_response.status_code, 303)
        detail_path = upload_response.headers["location"]

        detail_response = self.client.get(detail_path)
        self.assertEqual(detail_response.status_code, 200)
        self.assertIn("completed", detail_response.text)

        jobs_response = self.client.get("/jobs")
        self.assertEqual(jobs_response.status_code, 200)
        self.assertIn("history.pdf", jobs_response.text)

    def test_api_rejects_invalid_pdf(self):
        response = self.client.post(
            "/api/v1/process",
            headers={"Authorization": "Bearer token-123"},
            files=[("files", ("bad.pdf", b"not-a-pdf", "application/pdf"))],
        )
        self.assertEqual(response.status_code, 400)

    def test_api_rejects_when_concurrency_limit_reached(self):
        acquired = self.app.state.job_semaphore.acquire(blocking=False)
        self.assertTrue(acquired)
        try:
            response = self.client.post(
                "/api/v1/process",
                headers={"Authorization": "Bearer token-123"},
                files=[("files", self.fake_pdf_upload("busy.pdf"))],
            )
            self.assertEqual(response.status_code, 429)
        finally:
            self.app.state.job_semaphore.release()

    def test_templates_page_lists_current_templates(self):
        self.login()
        response = self.client.get("/templates")
        self.assertEqual(response.status_code, 200)
        self.assertIn("模板管理", response.text)
        self.assertIn("储蓄险单独总结书模板", response.text)
        self.assertTrue(self.app.state.template_store.current_path("savings_single").exists())
        self.assertNotIn("转换状态", response.text)

    def test_template_upload_and_restore(self):
        self.login()
        template_store = self.app.state.template_store
        current_path = template_store.current_path("savings_single")
        original_bytes = current_path.read_bytes()

        replacement_dir = Path(self.tempdir.name) / "replacement"
        replacement_dir.mkdir(parents=True, exist_ok=True)
        replacement_path = replacement_dir / "template_savings_standalone.docx"
        document = Document()
        document.add_paragraph("客户：{name}")
        document.add_paragraph("年龄：{age}")
        document.save(replacement_path)

        upload_response = self.client.post(
            "/templates/savings_single/upload",
            files={
                "file": (
                    "template_savings_standalone.docx",
                    replacement_path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            follow_redirects=False,
        )
        self.assertEqual(upload_response.status_code, 303)
        self.assertEqual(current_path.read_bytes(), replacement_path.read_bytes())

        template_data = template_store.get_template("savings_single")
        self.assertIn("{name}", template_data["placeholder_summary"])

        self.assertGreaterEqual(len(template_data["history"]), 1)
        version_name = template_data["history"][0]["name"]

        html_response = self.client.get("/templates/savings_single/html")
        self.assertEqual(html_response.status_code, 404)

        restore_response = self.client.post(
            "/templates/savings_single/restore",
            data={"version_name": version_name},
            follow_redirects=False,
        )
        self.assertEqual(restore_response.status_code, 303)
        self.assertEqual(current_path.read_bytes(), original_bytes)


if __name__ == "__main__":
    unittest.main()
